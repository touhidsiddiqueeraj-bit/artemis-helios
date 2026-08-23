"""
revision16.py — Re-embed corrected figures (GA, Figs 13-15)
============================================================
The figure PNGs were regenerated after visual review (GA redesign,
schematic restructure, waveform window fix, eff/loss title fix).
The docx still holds the OLD image binaries. This script finds the
four image paragraphs (column width 3187700 EMU, identified by their
extent aspect ratios) and replaces each drawing with the corrected PNG.

Run:  python3 backups/revision16.py
"""
from docx import Document
from docx.shared import Emu

DOC = '25195-52952-1-SM-REVISED.docx'
W_ = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
A_ = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
COL_W_EMU = 3187700

# identify by old extent cy (cx * old_h / old_w) — actual cy from the docx
TARGETS = {
    1487593: 'Figures/graphical_abstract.png',          # 1800x840
    1703070: 'Figures/fig13_power_stage_schematic.png', # actual cy
    2276929: 'Figures/fig15_buck_waveforms.png',         # 2100x1500
    1548311: 'Figures/fig16_eff_loss.png',              # actual cy
}

doc = Document(DOC)
body = doc.element.body
replaced = []

for p in body.iter(f'{W_}p'):
    ext = p.find(f'.//{A_}ext')
    if ext is None:
        continue
    if ext.get('cx') != str(COL_W_EMU):
        continue
    cy = int(ext.get('cy'))
    if cy not in TARGETS:
        continue
    path = TARGETS[cy]
    # clear all runs in this paragraph, keep pPr
    for r in p.findall(f'{W_}r'):
        p.remove(r)
    # wrap as python-docx paragraph to add picture
    from docx.text.paragraph import Paragraph
    para = Paragraph(p, doc)
    para.add_run().add_picture(path, width=Emu(COL_W_EMU))
    replaced.append(path)

assert len(replaced) == 4, f'expected 4 replacements, got {replaced}'
doc.save(DOC)
print('replaced:', *replaced, sep='\n  ')
