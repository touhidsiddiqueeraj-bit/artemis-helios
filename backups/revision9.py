"""
revision9.py — Round-9 layout & attribution fixes:
  1. Re-embed regenerated figures (blip-order mapping, same as revision8)
  2. Widen all figure images to full text width (7.16 in); QR stays 2.51 in
  3. Author Contributions: O.C. credited only for hardware validation
  4. Delete 15 empty paragraphs before ACKNOWLEDGEMENTS (page-8 whitespace)
  5. Verification asserts
"""
import re, shutil, os
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

DOCX = '25195-52952-1-SM-REVISED.docx'
BAK = '25195-52952-1-SM-REVISED-BAK9.docx'
NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
EMU_IN = 914400
FULL_W = int(7.16 * EMU_IN)

if not os.path.exists(BAK):
    shutil.copy2(DOCX, BAK)
    print(f"Backup -> {BAK}")

doc = Document(DOCX)
body = doc.element.body


def para_text(el):
    return ''.join(t.text or '' for t in el.iter() if t.tag.endswith('}t'))


def replace_run_text(el, new_text):
    for child in list(el):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('r', 'hyperlink'):
            el.remove(child)
    r = etree.SubElement(el, f'{{{NS_W}}}r')
    t = etree.SubElement(r, f'{{{NS_W}}}t')
    t.text = new_text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')


print("[1] Re-embed figures (blip order = display order)...")
FIG_SRC = {
    1:  'Code/Python/figures/fig1_architecture.png',
    2:  'Code/Python/figures/fig2_irradiance.png',
    3:  'Code/Python/figures/fig3_iv_curves.png',
    4:  'Code/Python/figures/fig4_lstm.png',
    5:  'Code/Python/figures/fig5_simulation.png',
    6:  'Code/Python/figures/fig6_comparison.png',
    7:  'Code/Python/figures/fig7_po_convergence.png',
    8:  'Code/Python/figures/fig9_validation.png',
    9:  'Logger_Data/cleaned/fig_validation_ramprates.png',
    10: 'Code/Python/figures/fig8_cost.png',
    11: 'Code/documentation/github_qr.png',
}
blip_paras = [el for el in body.iterchildren(f'{{{NS_W}}}p') if el.findall('.//' + qn('a:blip'))]
assert len(blip_paras) == 11, f"expected 11 blip paragraphs, got {len(blip_paras)}"
for n, el in enumerate(blip_paras, start=1):
    rid = el.findall('.//' + qn('a:blip'))[0].get(qn('r:embed'))
    part = doc.part.rels[rid].target_part
    with open(FIG_SRC[n], 'rb') as f:
        new_blob = f.read()
    if part.blob != new_blob:
        part._blob = new_blob
        print(f"  blip {n:2d} ({os.path.basename(FIG_SRC[n])}): replaced {len(new_blob)/1024:.0f} KB")
    else:
        print(f"  blip {n:2d}: unchanged")

print("[2] Widen figures to 7.16 in (QR stays)...")
from docx.oxml.ns import qn as _q
W = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
changed = 0
for i, el in enumerate(blip_paras):
    extent = el.findall('.//' + _q('wp:extent'))
    if not extent:
        continue
    ext = extent[0]
    cx = int(ext.get('cx'))
    if i == 10:  # QR (last blip) keep width
        print(f"  blip {i+1}: QR kept at {cx/EMU_IN:.2f} in")
        continue
    if cx == FULL_W:
        print(f"  blip {i+1}: already {cx/EMU_IN:.2f} in")
        continue
    scale = FULL_W / cx
    ext.set('cx', str(FULL_W))
    ext.set('cy', str(int(int(ext.get('cy')) * scale)))
    changed += 1
    print(f"  blip {i+1}: {cx/EMU_IN:.2f} -> {FULL_W/EMU_IN:.2f} in")
assert changed == 10, f"expected 10 width changes, got {changed}"

print("[3] Author Contributions rewrite...")
found = False
for el in body.iterchildren(f'{{{NS_W}}}p'):
    if para_text(el).strip() == 'Author Contributions Statement':
        nxt = el.getnext()
        replace_run_text(nxt,
            'Conceptualisation, H.T.S.; methodology, H.T.S.; software, H.T.S.; validation, '
            'H.T.S.; formal analysis, H.T.S.; investigation, H.T.S.; resources, H.T.S.; data '
            'curation, H.T.S.; hardware validation and field deployment, O.C.; '
            'writing—original draft preparation, H.T.S.; writing—review and editing, H.T.S.; '
            'supervision, H.T.S.; project administration, H.T.S.')
        found = True
        break
assert found, "Author Contributions heading not found"

print("[4] Delete empty paragraphs before ACKNOWLEDGEMENTS...")
els = list(body.iterchildren(f'{{{NS_W}}}p'))
ack_idx = next(i for i, e in enumerate(els) if para_text(e).strip() == 'ACKNOWLEDGEMENTS')
deleted = 0
for e in els[ack_idx - 20:ack_idx]:
    if not para_text(e).strip() and not e.findall('.//' + qn('a:blip')):
        body.remove(e)
        deleted += 1
print(f"  deleted {deleted} empty paragraphs")
assert deleted >= 10, f"expected >=10 empties, got {deleted}"

doc.save(DOCX)
print("saved")

print("[5] Verify...")
doc = Document(DOCX)
ps = [p.text for p in doc.paragraphs]
auth = next(p.text for p in doc.paragraphs if 'Conceptualisation' in p.text)
assert 'hardware validation and field deployment, O.C.' in auth
assert 'methodology, H.T.S.' in auth and 'data curation, H.T.S.' in auth
assert 'H.T.S. and O.C.' not in auth, "stale joint credit remains"
ack = next(i for i, p in enumerate(ps) if p.strip() == 'ACKNOWLEDGEMENTS')
empties = sum(1 for p in ps[ack-8:ack] if not p.strip())
print("  empties before ACKNOWLEDGEMENTS:", empties)
nrefs = sum(1 for p in ps if re.match(r'^\[\d+\]', p.strip()))
print("  refs:", nrefs)
assert nrefs == 37
wps = [e for e in doc.element.body.iterchildren(f'{{{NS_W}}}p') if e.findall('.//' + _q('a:blip'))]
widths = [int(e.findall('.//' + _q('wp:extent'))[0].get('cx')) / EMU_IN for e in wps]
print("  image widths (in):", [f"{w:.2f}" for w in widths])
assert all(abs(w - 7.16) < 0.01 for w in widths[:-1]) and abs(widths[-1] - 2.51) < 0.01
print("ALL CHECKS PASSED ✓")
