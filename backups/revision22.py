"""
revision22.py — v3 layout fixes: Fig13 re-embed, Table IV rebuild, section reorder
==================================================================================
1) Re-embed regenerated field-day figure (Fig. 13).
2) Rebuild Table IV: fixed layout, column width 5020 twips, 7 pt font, short labels.
3) Move the sensitivity heatmap (descriptive paragraph + Fig. 14 + caption) from
   V.A into a new Section IV.I "Sensitivity Analysis" (results belong with results).
4) Swap sections V <-> VI: hardware implementation becomes V, discussion becomes VI
   (discussion no longer precedes hardware). Figure re-serialization: logger fig
   16 -> 15, BOM fig 15 -> 16; section refs updated.

Run:  python3 backups/revision22.py
"""
import re
import copy
from lxml import etree
from docx import Document
from docx.shared import Emu, Pt

DOC = '25195-52952-1-SM-REVISED.docx'
W_ = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
A_ = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
R_ = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'
COL_W_EMU = 3187700
COL_W_TWIPS = 5020
FIG = 'Figures/fig_field_day_evidence.png'

doc = Document(DOC)
body = doc.element.body
paras = list(doc.paragraphs)

def ptext(p):
    return ''.join(r.text or '' for r in p.runs)

def set_text(p, new):
    for r in p.runs:
        r.text = ''
    p.runs[0].text = new

# ── 1) re-embed Fig. 13 (field-day figure) ──────────────────────────────────
cap13 = next(p for p in paras if ptext(p).strip().startswith('Fig. 13. Efficiency evidence'))
img13 = cap13._p.getprevious()
assert img13 is not None and img13.tag == f'{W_}p' and img13.find(f'.//{A_}blip') is not None
for r in img13.findall(f'{W_}r'):
    img13.remove(r)
from docx.text.paragraph import Paragraph
Paragraph(img13, doc).add_run().add_picture(FIG, width=Emu(COL_W_EMU))
print('fig13 re-embedded')

# ── 2) rebuild Table IV ─────────────────────────────────────────────────────
tbl_el = None
for t in doc.tables:
    if t.cell(0, 0).text.strip().startswith('Scenario'):
        tbl_el = t._tbl
        break
assert tbl_el is not None
tbl_el.getparent().remove(tbl_el)

def build_table():
    tbl = etree.Element(f'{W_}tbl')
    tblPr = etree.SubElement(tbl, f'{W_}tblPr')
    tblW = etree.SubElement(tblPr, f'{W_}tblW')
    tblW.set(W_ + 'w', str(COL_W_TWIPS)); tblW.set(W_ + 'type', 'dxa')
    etree.SubElement(tblPr, f'{W_}tblLayout').set(W_ + 'type', 'fixed')
    etree.SubElement(tblPr, f'{W_}jc').set(W_ + 'val', 'center')
    borders = etree.SubElement(tblPr, f'{W_}tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = etree.SubElement(borders, f'{W_}{side}')
        b.set(W_ + 'val', 'single'); b.set(W_ + 'sz', '4'); b.set(W_ + 'color', '000000')
    widths = [2100, 1000, 1000, 920]
    grid = etree.SubElement(tbl, f'{W_}tblGrid')
    for wd in widths:
        etree.SubElement(grid, f'{W_}gridCol').set(W_ + 'w', str(wd))
    data = [
        ('Scenario', 'P&O (%)', 'LSTM (%)', 'Gap (pp)'),
        ('Whole day', '81.9', '93.3', '11.4'),
        ('High-variability 20%', '67.8', '90.2', '22.4'),
        ('Low-variability 80%', '84.9', '94.0', '9.1'),
        ('Ramp < 10 W/m\u00b2/min', '84.9', '94.0', '9.1'),
        ('Ramp 10\u201350 W/m\u00b2/min', '84.9', '94.0', '9.1'),
        ('Ramp 50\u2013150 W/m\u00b2/min', '84.8', '93.7', '8.9'),
        ('Ramp > 150 W/m\u00b2/min', '51.1', '86.9', '35.8'),
    ]
    for ri, row in enumerate(data):
        tr = etree.SubElement(tbl, f'{W_}tr')
        for ci, val in enumerate(row):
            tc = etree.SubElement(tr, f'{W_}tc')
            tcPr = etree.SubElement(tc, f'{W_}tcPr')
            tcW = etree.SubElement(tcPr, f'{W_}tcW')
            tcW.set(W_ + 'w', str(widths[ci])); tcW.set(W_ + 'type', 'dxa')
            etree.SubElement(tcPr, f'{W_}vAlign').set(W_ + 'val', 'center')
            p = etree.SubElement(tc, f'{W_}p')
            r = etree.SubElement(p, f'{W_}r')
            rPr = etree.SubElement(r, f'{W_}rPr')
            sz = etree.SubElement(rPr, f'{W_}sz'); sz.set(W_ + 'val', '14')  # 7 pt
            if ri == 0:
                b = etree.SubElement(rPr, f'{W_}b')
            t = etree.SubElement(r, f'{W_}t'); t.set(W_ + 'space', 'preserve')
            t.text = val
    return tbl

new_tbl = build_table()
tcap = next(p for p in paras if ptext(p).strip().startswith('Table IV.'))
tcap._p.addnext(new_tbl)
print('table IV rebuilt')

# ── 3) section surgery ──────────────────────────────────────────────────────
disc_h = next(p for p in paras if ptext(p).strip().startswith('V.  DISCUSSION'))
hw_h   = next(p for p in paras if ptext(p).strip().startswith('VI. HARDWARE IMPLEMENTATION'))
vii_h  = next(p for p in paras if ptext(p).strip().startswith('VII.  CONCLUSION'))
cap14 = next(p for p in paras if ptext(p).strip().startswith('Fig. 14. Multidimensional'))
img14 = cap14._p.getprevious()
assert img14 is not None and img14.tag == f'{W_}p' and img14.find(f'.//{A_}blip') is not None
para104 = img14.getprevious()
assert para104 is not None and para104.tag == f'{W_}p' and 'This analysis extends' in ptext(
    Paragraph(para104, doc))
print('found heatmap block + section anchors')

# collect hardware block (hw_h .. before vii_h)
hw_block = []
el = hw_h._p
while el is not None and el != vii_h._p:
    hw_block.append(el)
    el = el.getnext()

# detach heatmap block and hardware block
for el in (para104, img14, cap14._p):
    el.getparent().remove(el)
for el in hw_block:
    el.getparent().remove(el)

# retitle headings
set_text(hw_h, 'V. HARDWARE IMPLEMENTATION')
set_text(disc_h, 'VI.  DISCUSSION')

# IV.I heading (clone of IV.H heading style)
subh = next(p for p in paras if ptext(p).strip().startswith('H.  Efficiency Evidence'))
i_heading = copy.deepcopy(subh._p)
runs = i_heading.findall(f'{W_}r')
first = runs[0]
for r in runs[1:]:
    i_heading.remove(r)
ts = first.findall(f'{W_}t')
ts[0].text = 'I.  Sensitivity Analysis'
for t in ts[1:]:
    first.remove(t)

# reinsert before disc heading: IV.I heading, heatmap block, hardware block
anchor = disc_h._p
for el in ([i_heading, para104, img14, cap14._p] + hw_block):
    anchor.addprevious(el)
print('sections reordered: IV.I (sensitivity) + V hardware before VI discussion')

# ── 4) ref remaps ───────────────────────────────────────────────────────────
ref_fixes = [
    ('presented in Section VI (Fig. 16)', 'presented in Section V (Fig. 15)'),
    ('Fig. 15presents the component cost breakdown', 'Fig. 16 presents the component cost breakdown'),
    ('Fig. 15.  BOM breakdown', 'Fig. 16.  BOM breakdown'),
    ('data-acquisition board (Fig. 16)', 'data-acquisition board (Fig. 15)'),
    ('Fig. 16. Field logger schematic', 'Fig. 15. Field logger schematic'),
]
n = 0
for p in paras:
    t = ptext(p)
    nt = t
    for old, new in ref_fixes:
        if old in nt:
            nt = nt.replace(old, new)
    nt = re.sub(r'\b(Fig\.\s*\d+)(presents|shows|gives)', r'\1 \2', nt)
    if nt != t:
        set_text(p, nt)
        n += 1
print('ref fixes applied to', n, 'paragraphs')

doc.save(DOC)
print('revision22 saved')
