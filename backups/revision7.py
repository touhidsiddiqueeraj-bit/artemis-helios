"""
revision7.py — Fix all remaining issues (v2: run-based, no index fragility)
"""
import re, shutil, os
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

DOCX = '25195-52952-1-SM-REVISED.docx'
BAK = '25195-52952-1-SM-REVISED-BAK7.docx'
NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

if not os.path.exists(BAK):
    shutil.copy2(DOCX, BAK)
    print(f"Backup -> {BAK}")

doc = Document(DOCX)
body = doc.element.body

def replace_run_text(el, new_text):
    """Replace all run text in an element, preserving structure."""
    for child in list(el):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('r', 'hyperlink'):
            el.remove(child)
    r = etree.SubElement(el, f'{{{NS_W}}}r')
    t = etree.SubElement(r, f'{{{NS_W}}}t')
    t.text = new_text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

def replace_in_runs(el, old_str, new_str):
    """Replace text within each run of a paragraph element."""
    for child in el.iter():
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 't' and child.text and old_str in child.text:
            child.text = child.text.replace(old_str, new_str)

def para_el_text(el):
    """Reconstruct text from an element's runs."""
    parts = []
    for child in el.iter():
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 't' and child.text:
            parts.append(child.text)
    return ''.join(parts)

# Get all paragraph elements for stable references
all_p_els = list(body.iterchildren(f'{{{NS_W}}}p'))
print(f"Total paragraph elements: {len(all_p_els)}")

def find_el_idx(text_fragment):
    for idx, el in enumerate(all_p_els):
        if text_fragment in para_el_text(el):
            return idx, el
    return None, None

# ─── 1. ADD CONTENT TO I.C ────────────────────────────────────────
print("[1] I.C — add content...")
ic_idx, ic_el = find_el_idx('C.  State of the Art')
if ic_idx is not None:
    next_el = all_p_els[ic_idx + 1] if ic_idx + 1 < len(all_p_els) else None
    if next_el and para_el_text(next_el).strip().startswith('D.  '):
        new_p = etree.SubElement(body, f'{{{NS_W}}}p')
        body.insert(list(body).index(ic_el) + 1, new_p)
        replace_run_text(new_p, (
            'Intelligent MPPT techniques have evolved along two principal axes: '
            'machine-learning-based irradiance forecasting for predictive control, '
            'and adaptive hill-climbing methods that improve upon fixed-step P&O. '
            'On the forecasting axis, LSTM networks [23],[8],[9] have demonstrated '
            'superior ability to model the non-linear temporal dynamics of solar '
            'irradiance compared with feed-forward or convolutional architectures, '
            'particularly under the high-frequency cloud flicker characteristic of '
            'tropical monsoon climates. On the adaptive axis, variable-step P&O [15] '
            'and incremental conductance [1] reduce steady-state oscillation but '
            'remain fundamentally reactive. The dual-MCU paradigm introduced in this '
            'work occupies the intersection of both axes.'
        ))
        all_p_els.insert(ic_idx + 1, new_p)
        print(f"  I.C content inserted ✓")

# ─── 2. ADD CONTENT TO I.E ────────────────────────────────────────
print("[2] I.E — add content...")
ie_idx, ie_el = find_el_idx('E.  Contribution')
if ie_idx is not None:
    next_el = all_p_els[ie_idx + 1] if ie_idx + 1 < len(all_p_els) else None
    if next_el and para_el_text(next_el).strip().startswith('F.  '):
        new_p = etree.SubElement(body, f'{{{NS_W}}}p')
        body.insert(list(body).index(ie_el) + 1, new_p)
        replace_run_text(new_p, (
            'This paper presents the Helios-Artemis dual-MCU predictive MPPT '
            'controller, which partitions the problem into an ESP32-S3 prediction '
            'plane (LSTM irradiance forecasting, local retraining, SD logging) '
            'and an STM32F103 control plane (real-time VS-P&O, PWM generation, '
            'battery management). The key contributions are: (1) an LSTM-assisted '
            'MPPT that achieves 94.0% mean monsoon tracking efficiency vs 70.7% '
            'for plain P&O in simulation; (2) a zero-cloud architecture avoiding '
            'the connectivity dependency of prior LSTM-MPPT designs; (3) a '
            'Markov+OU synthetic irradiance model validated against field '
            'measurements; and (4) a sub-USD 17 BOM compatible with IDCOL SHS '
            'cost targets.'
        ))
        all_p_els.insert(ie_idx + 1, new_p)
        print(f"  I.E content inserted ✓")

# ─── 3. FIX FIG. 10 COST TEXT ─────────────────────────────────────
print("[3] Cost text — add 250 BDT assembly...")
cost_idx, cost_el = find_el_idx('Fig. 10 presents the component cost breakdown')
if cost_idx is not None:
    new_text = (
        'Fig. 10 presents the component cost breakdown. The estimated controller '
        'component cost is approximately 1,750 BDT (USD 16), comprising ESP32-S3 '
        '(380 BDT), STM32F103 (120 BDT), INA219 (80 BDT), TSL2591 (120 BDT), buck '
        'converter passives (350 BDT), PCB and housing (280 BDT), assembly (250 BDT), '
        'and miscellaneous connectors (170 BDT). This represents an 87% reduction '
        'versus IDCOL-compatible commercial MPPT controllers (13,500 BDT). At a rural '
        'avoided cost of 60 BDT/kWh, the projected annual yield improvement of +4.8 '
        'kWh per 50 Wp panel (91.3% vs 85.8% weighted annual efficiency, from Table '
        'III Monte Carlo annual means) generates +289 BDT/year additional energy value, '
        'implying a payback period of approximately 6.1 years against a plain P&O '
        'controller baseline.'
    )
    replace_run_text(cost_el, new_text)
    print(f"  Cost text updated ✓")

# ─── 4. FIX ACKNOWLEDGEMENTS ─────────────────────────────────────
print("[4] Acknowledgements — SUST -> Leading University...")
ack_idx, ack_el = find_el_idx('Shahjalal University of Science and Technology, Sylhet')
if ack_idx is not None:
    for child in ack_el.iter():
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 't' and child.text and 'Shahjalal University' in child.text:
            child.text = child.text.replace(
                'Shahjalal University of Science and Technology, Sylhet',
                'Leading University, Sylhet'
            )
    print(f"  Acknowledgements updated ✓")

# ─── 5. FIX AUTHOR CONTRIBUTIONS ─────────────────────────────────
print("[5] Author Contributions...")
ac_idx, ac_el = find_el_idx('Author Contributions Statement')
if ac_idx is not None:
    next_el = all_p_els[ac_idx + 1]
    if next_el is not None:
        replace_run_text(next_el, (
            'Conceptualisation, H.T.S.; methodology, H.T.S. and O.C.; '
            'software, H.T.S.; validation, H.T.S. and O.C.; formal analysis, '
            'H.T.S.; investigation, H.T.S. and O.C.; resources, H.T.S.; data '
            'curation, O.C.; hardware deployment, O.C.; writing\u2014original draft '
            'preparation, H.T.S.; writing\u2014review and editing, H.T.S. and O.C.; '
            'supervision, H.T.S.; project administration, H.T.S.'
        ))
        print(f"  Author Contributions updated ✓")

# ─── 6. REMOVE AUTHOR BIOGRAPHY ──────────────────────────────────
print("[6] Remove Author Biography...")
for target in ['Orpon Chanda received the B.Sc.', 'Hussain Touhid Siddiquee received', 'Author Biography']:
    for el in list(body.iterchildren(f'{{{NS_W}}}p')):
        text = para_el_text(el).strip()
        if text.startswith(target):
            body.remove(el)
            all_p_els[:] = [e for e in all_p_els if e != el]
            print(f"  Removed: '{text[:50]}...' ✓")
            break

# refresh all_p_els
all_p_els = list(body.iterchildren(f'{{{NS_W}}}p'))

# ─── 7. RE-EMBED FIGURES ──────────────────────────────────────────
print("[7] Re-embed figures...")
fig_src = os.path.join(os.path.dirname(__file__) or '.', 'Code/Python/figures')
for fname in ['fig1_architecture.png', 'fig8_cost.png']:
    fpath = os.path.join(fig_src, fname)
    if os.path.exists(fpath):
        # Find the image part in docx
        for p_el in body.iterchildren(f'{{{NS_W}}}p'):
            blips = p_el.findall('.//' + qn('a:blip'))
            if blips:
                rid = blips[0].get(qn('r:embed'))
                if rid and rid in doc.part.rels:
                    target = doc.part.rels[rid].target_part
                    with open(fpath, 'rb') as f:
                        target._blob = f.read()
                    print(f"  {fname} -> {target.partname} ({len(target._blob)} bytes)")
                    break
        else:
            print(f"  SKIP: no image element for {fname}")
    else:
        print(f"  SKIP: {fpath} not found")

# ─── 8. REFERENCE DEDUP + RENUMBER ────────────────────────────────
print("[8] Reference dedup and renumbering...")

# Step 8a: Find REFERENCES heading
ref_heading_idx = None
for idx, el in enumerate(all_p_els):
    if para_el_text(el).strip() == 'REFERENCES':
        ref_heading_idx = idx
        break

if ref_heading_idx is not None:
    # Collect ref list paragraphs
    ref_els = []
    for el in all_p_els[ref_heading_idx + 1:]:
        text = para_el_text(el).strip()
        if not re.match(r'^\[\d+\]', text):
            break
        ref_els.append(el)

    print(f"  Found {len(ref_els)} reference entries")

    # Dedup: [29] = [23], [40] = [35]
    # Apply dedup in ALL runs first
    dedup_map = {29: 23, 40: 35}
    for old_n, new_n in dedup_map.items():
        old_str = f'[{old_n}]'
        new_str = f'[{new_n}]'
        for el in all_p_els:
            replace_in_runs(el, old_str, new_str)
    print(f"  Dedup applied: {dedup_map}")

    # Remove [29] and [40] from ref list
    for el in ref_els[:]:
        text = para_el_text(el)
        m = re.match(r'\[(\d+)\]', text)
        if m and int(m.group(1)) in dedup_map:
            body.remove(el)
            ref_els.remove(el)
            print(f"  Removed duplicate: {text[:60]}...")

    # Refresh all_p_els
    all_p_els = list(body.iterchildren(f'{{{NS_W}}}p'))

    # Step 8b: Determine citation order
    seen = set()
    order = []
    for el in all_p_els:
        text = para_el_text(el)
        for m in re.finditer(r'\[(\d+)\]', text):
            n = int(m.group(1))
            if n not in seen:
                seen.add(n)
                order.append(n)

    print(f"  Citation order: {order}")
    remap = {old: i+1 for i, old in enumerate(order)}
    print(f"  Remap: {remap}")

    # Step 8c: Apply renumbering to ALL runs
    for old_n in sorted(remap.keys(), reverse=True):
        new_n = remap[old_n]
        old_str = f'[{old_n}]'
        new_str = f'[{new_n}]'
        for el in all_p_els:
            replace_in_runs(el, old_str, new_str)

    # Step 8d: Update ref list numbering and reorder
    # Re-read ref entries after all text changes
    new_ref_els = []
    for el in list(body.iterchildren(f'{{{NS_W}}}p')):
        text = para_el_text(el).strip()
        m = re.match(r'^\[(\d+)\]', text)
        if not m:
            continue
        # Check if it's in the ref section (after REFERENCES)
        # Simple: use the el's position relative to ref_heading
        n = int(m.group(1))
        new_ref_els.append((n, el, text))

    # Remove all from body
    ref_parent = ref_heading_idx  # pylance ignore
    heading_el = all_p_els[ref_heading_idx]
    for _, el, _ in new_ref_els:
        try:
            body.remove(el)
        except ValueError:
            pass

    # Sort and reinsert
    new_ref_els.sort(key=lambda x: x[0])
    for idx, (n, el, text) in enumerate(new_ref_els):
        body.insert(list(body).index(heading_el) + 1 + idx, el)

    print(f"  Reference list reordered ✓ ({len(new_ref_els)} refs)")

# ─── 9. FONTS ─────────────────────────────────────────────────────
print("[9] IJPEDS style formatting...")
for el in all_p_els:
    for rPr in el.findall(f'.//{{{NS_W}}}rPr'):
        rFonts = rPr.find(f'{{{NS_W}}}rFonts')
        if rFonts is None:
            rFonts = etree.SubElement(rPr, f'{{{NS_W}}}rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Times New Roman')
print("  Times New Roman set ✓")

# ─── 10. SAVE ─────────────────────────────────────────────────────
print("[10] Saving...")
doc.save(DOCX)
print("Saved ✓")

# ─── 11. VERIFICATION ─────────────────────────────────────────────
print("\n[11] Verification:")
doc2 = Document(DOCX)
checks = [
    ('I.C has content', lambda: any('Intelligent MPPT' in p.text for p in doc2.paragraphs)),
    ('I.E has content', lambda: any('Helios-Artemis dual-MCU' in p.text for p in doc2.paragraphs)),
    ('Cost has assembly', lambda: any('assembly (250 BDT)' in p.text for p in doc2.paragraphs)),
    ('Cost total 1,750', lambda: any('1,750 BDT' in p.text for p in doc2.paragraphs)),
    ('Cost 87% reduction', lambda: any('87% reduction' in p.text for p in doc2.paragraphs)),
    ('Ack: Leading University', lambda: any('Leading University, Sylhet' in p.text for p in doc2.paragraphs)),
    ('Ack: no SUST', lambda: not any('Shahjalal University of Science and Technology, Sylhet' in p.text for p in doc2.paragraphs)),
    ('Contrib: hardware deployment OC', lambda: any('hardware deployment, O.C.' in p.text for p in doc2.paragraphs)),
    ('No Author Biography text', lambda: not any('Orpon Chanda received' in p.text for p in doc2.paragraphs)),
    ('No [29] in body', lambda: not any('[29]' in p.text for p in doc2.paragraphs)),
    ('No [40] in body', lambda: not any('[40]' in p.text for p in doc2.paragraphs)),
    ('No duplicate ref numbers', lambda: (
        len(set(re.findall(r'^\[(\d+)\]', p.text.strip()) for p in doc2.paragraphs if re.match(r'^\[\d+\]', p.text.strip())))
        == len([re.match(r'^\[\d+\]', p.text.strip()).group(1) for p in doc2.paragraphs if re.match(r'^\[\d+\]', p.text.strip())])
    )),
]
all_ok = True
for label, fn in checks:
    ok = fn()
    print(f'  {"✓" if ok else "✗"} {label}')
    if not ok: all_ok = False

# Check citation sequential order
import re as re_m
seen = set()
last_ok = True
for p in doc2.paragraphs:
    for r in re_m.findall(r'\[(\d+)\]', p.text):
        n = int(r)
        if n not in seen:
            seen.add(n)
            if n != len(seen):
                print(f'  ✗ Citation order broken: expected [{len(seen)}], found [{n}] in "{p.text[:50]}..."')
                last_ok = False
                break
    if not last_ok:
        break
if last_ok:
    print(f'  ✓ Citation order: sequential 1\u2013{len(seen)}')

print(f'\n{"All checks passed! ✓" if all_ok and last_ok else "Some checks failed ✗"}')
