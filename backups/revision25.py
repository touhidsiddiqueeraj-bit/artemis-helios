"""
revision25.py — one-dimensional sensitivity sweeps (new Fig. 15, renumber 15..19 -> 16..20)
============================================================================================
1) Renumber figures 15..19 -> 16..20 (transient suite, field logger,
   timing budget, BOM, QR) in captions and in-text refs.
2) Insert after the Fig. 14 caption paragraph (end of the sensitivity
   subsection, immediately before the IV.J heading):
   - paragraph reporting the three sweeps (forecast window, P&O step,
     UART-latency stand-in) with headline numbers
   - Fig. 15 (Figures/fig15_sensitivity_sweeps.png) with caption

Data: Code/Python/results/sensitivity_sweeps.csv
Run:  python3 backups/revision25.py
"""
import re
import copy
import os
from lxml import etree
from docx import Document
from docx.shared import Emu
from docx.text.paragraph import Paragraph
from docx.oxml import parse_xml

DOC = '25195-52952-1-SM-REVISED.docx'
W_ = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
COL_W_EMU = 3187700
FIG = 'Figures/fig15_sensitivity_sweeps.png'

doc = Document(DOC)
body = doc.element.body
paras = list(doc.paragraphs)


def ptext(p):
    return ''.join(r.text or '' for r in p.runs)


def set_text(p, new):
    for r in p.runs:
        r.text = ''
    p.runs[0].text = new


def replace_run_text(el, text):
    for r in el.findall(f'{W_}r'):
        el.remove(r)
    run = etree.SubElement(el, f'{W_}r')
    rPr = etree.SubElement(run, f'{W_}rPr')
    sz = etree.SubElement(rPr, f'{W_}sz'); sz.set(W_ + 'val', '20')
    szcs = etree.SubElement(rPr, f'{W_}szCs'); szcs.set(W_ + 'val', '20')
    t = etree.SubElement(run, f'{W_}t'); t.set(W_ + 'space', 'preserve')
    t.text = text
    return el


def image_para():
    ns = W_.strip('{}')
    xml = (f'<w:p xmlns:w="{ns}">'
           f'<w:pPr><w:jc w:val="center"/></w:pPr>'
           f'<w:r/></w:p>')
    return parse_xml(xml)


def insert_after(anchor_el, new_el):
    anchor_el.addnext(new_el)


# ── 1) renumber Fig. 15..19 -> 16..20 (one pass, no collisions) ─────────────
renum = re.compile(r'Fig\. (1[5-9])')
n = 0
for p in paras:
    t = ptext(p)
    nt = renum.sub(lambda m: 'Fig. ' + str(int(m.group(1)) + 1), t)
    if nt != t:
        set_text(p, nt)
        n += 1
print('renumbered Fig. refs in', n, 'paragraphs')

# ── 2) insertion after the Fig. 14 caption paragraph ───────────────────────
fig14_cap = next(p for p in doc.paragraphs if ptext(p).strip().startswith(
    'Fig. 14. Multidimensional sensitivity'))
anchor = fig14_cap._p

# body paragraph cloned from the Fig. 14 caption (Normal, 10 pt)
tmpl = fig14_cap

INTRO = ('A one-dimensional sweep completes the sensitivity picture for '
         'the three remaining design dimensions \u2014 forecast-memory '
         'window, P&O step size and control-loop (UART) latency \u2014 on '
         'the same stochastic day (seed 23, dt = 0.1 s, Fig. 15). Tracking '
         'efficiency is insensitive to the forecast window: 94.8\u201395.3% '
         'across 1\u201360 s with a shallow optimum at 10 s, so the blend '
         'neither needs nor suffers from long prediction memory on '
         'sub-minute dynamics. Efficiency favours smaller P&O steps '
         '(97.1% at 0.1 V versus 95.3\u201395.5% at 0.8\u20131.6 V), '
         'consistent with the variable-step design adapting its step to the '
         'irradiance gradient. Loop latency costs at most two percentage '
         'points for delays up to 2 s (95.3% at zero delay versus 93.3% at '
         '0.5 s, recovering to 94.0% at 1 s), and the measured 3.48 ms '
         'Helios\u2013Artemis link (Table V) sits three orders of magnitude '
         'below the onset of any degradation.')

CAPTION = ('Fig. 15. One-dimensional sensitivity of LSTM-P&O tracking '
           'efficiency on the stochastic day (seed 23, dt = 0.1 s): '
           '(a) forecast-memory window, (b) variable-step P&O size '
           'delta_max, (c) control-loop (UART) latency in 0.1 s ticks. '
           'Baselines: 10 s window, 0.80 V, zero delay.')

body_el = copy.deepcopy(tmpl._p)
replace_run_text(body_el, INTRO)
insert_after(anchor, body_el)

imgp = image_para()
insert_after(body_el, imgp)
Paragraph(imgp, doc).add_run().add_picture(FIG, width=Emu(COL_W_EMU))

cap = copy.deepcopy(tmpl._p)
replace_run_text(cap, CAPTION)
insert_after(imgp, cap)

doc.save(DOC)
print('Fig. 15 sweeps block inserted after Fig. 14 caption')
