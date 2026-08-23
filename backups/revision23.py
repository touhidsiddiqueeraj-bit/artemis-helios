"""
revision23.py — measured Helios execution-time results (Table V, new Fig. 16)
==============================================================================
1) Insert at the end of Section V (HARDWARE IMPLEMENTATION):
   - a measurement paragraph (ESP32-S3 probe, N = 400, 240 MHz)
   - Table V (measured timing budget, fixed layout, 7 pt)
   - Fig. 16 (timing budget + loop jitter) with caption
2) Renumber: BOM figure 16 -> 17 (in-text ref + caption), QR 17 -> 18.
3) Limitations (VI.C): note STM32-side latency pending hardware.
4) Future work (VI.C): add full closed-loop latency measurement.

Run:  python3 backups/revision23.py
"""
import re
import copy
from lxml import etree
from docx import Document
from docx.shared import Emu, Pt
from docx.text.paragraph import Paragraph
from docx.oxml import parse_xml

DOC = '25195-52952-1-SM-REVISED.docx'
W_ = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
A_ = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
COL_W_EMU = 3187700
COL_W_TWIPS = 5020
FIG = 'Figures/fig_timing_budget.png'

doc = Document(DOC)
body = doc.element.body
paras = list(doc.paragraphs)


def ptext(p):
    return ''.join(r.text or '' for r in p.runs)


def set_text(p, new):
    for r in p.runs:
        r.text = ''
    p.runs[0].text = new


# ── 1) timing block at the end of Section V ────────────────────────────────
silks = next(p for p in paras if ptext(p).strip().startswith(
    'Silkscreen markers on the PCB'))
hw_end = silks._p
disc_h = next(p for p in paras if ptext(p).strip().startswith('VI.  DISCUSSION'))
assert hw_end.getnext() == disc_h._p, 'silkscreen paragraph must precede VI heading'

vis_p = next(p for p in paras if ptext(p).strip().startswith(
    'The field irradiance logger employed'))
vis_el = vis_p._p
body_p = next(p for p in paras if ptext(p).strip().startswith(
    'To validate the Markov+OU'))


def clone_body_para(text):
    el = copy.deepcopy(body_p._p)
    for r in el.findall(f'{W_}r'):
        el.remove(r)
    run = etree.SubElement(el, f'{W_}r')
    rPr = etree.SubElement(run, f'{W_}rPr')
    sz = etree.SubElement(rPr, f'{W_}sz'); sz.set(W_ + 'val', '20')
    szcs = etree.SubElement(rPr, f'{W_}szCs'); szcs.set(W_ + 'val', '20')
    t = etree.SubElement(run, f'{W_}t'); t.set(W_ + 'space', 'preserve')
    t.text = text
    return el


def image_para():
    ns = W_.strip('{}')
    xml = (f'<w:p xmlns:w="{ns}">'
           f'<w:pPr><w:jc w:val="center"/></w:pPr>'
           f'<w:r/></w:p>')
    return parse_xml(xml)


def caption_para(text):
    el = copy.deepcopy(silks._p)
    for r in el.findall(f'{W_}r'):
        el.remove(r)
    run = etree.SubElement(el, f'{W_}r')
    rPr = etree.SubElement(run, f'{W_}rPr')
    b = etree.SubElement(rPr, f'{W_}b')
    sz = etree.SubElement(rPr, f'{W_}sz'); sz.set(W_ + 'val', '20')
    szcs = etree.SubElement(rPr, f'{W_}szCs'); szcs.set(W_ + 'val', '20')
    t = etree.SubElement(run, f'{W_}t'); t.set(W_ + 'space', 'preserve')
    t.text = text
    return el


def build_table():
    tbl = etree.Element(f'{W_}tbl')
    tblPr = etree.SubElement(tbl, f'{W_}tblPr')
    tblW = etree.SubElement(tblPr, f'{W_}tblW')
    tblW.set(W_ + 'w', str(COL_W_TWIPS)); tblW.set(W_ + 'type', 'dxa')
    etree.SubElement(tblPr, f'{W_}tblLayout').set(W_ + 'type', 'fixed')
    etree.SubElement(tblPr, f'{W_}jc').set(W_ + 'val', 'center')
    borders = etree.SubElement(tblPr, f'{W_}tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = etree.SubElement(borders, f'{W_}{side}')
        b.set(W_ + 'val', 'single'); b.set(W_ + 'sz', '4'); b.set(W_ + 'color', '000000')
    widths = [2700, 800, 760, 760]
    grid = etree.SubElement(tbl, f'{W_}tblGrid')
    for wd in widths:
        etree.SubElement(grid, f'{W_}gridCol').set(W_ + 'w', str(wd))
    data = [
        ('Stage', 'Mean', 'p99', 'Max'),
        ('Preprocessing (24-step window)', '7.1 \u00b5s', '8 \u00b5s', '\u2013'),
        ('LSTM inference (24 \u00d7 32 units)', '6.355 ms', '6.360 ms', '6.457 ms'),
        ('Packet formatting', '86.7 \u00b5s', '\u2013', '382 \u00b5s'),
        ('UART transmission (115.2 kbaud)', '3.484 ms', '\u2013', '3.491 ms'),
        ('Full Helios control tick', '9.996 ms', '\u2013', '\u2013'),
        ('Loop period (100 ms nominal)', '100.000 ms', '100.026 ms', '100.032 ms'),
        ('Absolute loop jitter', '16.3 \u00b5s', '31 \u00b5s', '36 \u00b5s'),
    ]
    for ri, row in enumerate(data):
        tr = etree.SubElement(tbl, f'{W_}tr')
        for ci, val in enumerate(row):
            tc = etree.SubElement(tr, f'{W_}tc')
            tcPr = etree.SubElement(tc, f'{W_}tcPr')
            tcW = etree.SubElement(tcPr, f'{W_}tcW')
            tcW.set(W_ + 'w', str(widths[ci])); tcW.set(W_ + 'type', 'dxa')
            etree.SubElement(tcPr, f'{W_}vAlign').set(W_ + 'val', 'center')
            p = etree.SubElement(tc, f'{W_}p')
            r = etree.SubElement(p, f'{W_}r')
            rPr = etree.SubElement(r, f'{W_}rPr')
            sz = etree.SubElement(rPr, f'{W_}sz'); sz.set(W_ + 'val', '14')
            if ri == 0:
                etree.SubElement(rPr, f'{W_}b')
            t = etree.SubElement(r, f'{W_}t'); t.set(W_ + 'space', 'preserve')
            t.text = val
    return tbl


MEAS_PAR = ('The Helios-side control budget was measured directly on the '
            'ESP32-S3 module (N16R8, 240 MHz) with a probe firmware that '
            'executes the paper-sized network (32 hidden units, 24-step '
            'lookback, 33 inputs) and the 115.2 kbaud UART packet link. '
            'Timestamps come from the ESP32 hardware timer and are emitted '
            'only after each batch of 400 runs, so serial output does not '
            'perturb the timed sections. The complete Helios tick '
            '(preprocessing, 24-step inference, packet formatting and UART '
            'output) averages 9.996 ms, one tenth of the 100 ms control '
            'cycle, and the loop period measured over 400 consecutive cycles '
            'stays within 32 \u00b5s of 100.000 ms at the maximum (Table V, '
            'Fig. 16). LSTM inference dominates the tick at 6.355 ms on '
            'average (p99 6.360 ms), leaving roughly 90 ms of idle time per '
            'cycle on the Helios core.')

CAPTION = ('Fig. 16. Measured Helios execution-time budget on the ESP32-S3 '
           '(N = 400, 240 MHz): (a) stacked control tick of 9.996 ms against '
           'the 100 ms cycle, (b) distribution of the measured 100 ms loop '
           'periods with the mean, p99 and maximum marked. Values are '
           'summarised in Table V.')

new_par = clone_body_para(MEAS_PAR)
tcap = caption_para('Table V. Measured Helios execution-time budget (ESP32-S3 '
                    'N16R8, 240 MHz, N = 400).')
tbl = build_table()
fig_p = image_para()
fcap = caption_para(CAPTION)

anchor = disc_h._p
for el in (new_par, tcap, tbl):
    anchor.addprevious(el)

Paragraph(fig_p, doc).add_run().add_picture(FIG, width=Emu(COL_W_EMU))
anchor.addprevious(fig_p)
anchor.addprevious(fcap)
print('timing block inserted (paragraph, Table V, Fig. 16)')

# ── 2) renumber 16 -> 17 (BOM) and 17 -> 18 (QR) ───────────────────────────
ref_fixes = [
    ('Fig. 16 presents the component cost breakdown',
     'Fig. 17 presents the component cost breakdown'),
    ('Fig. 16.  BOM breakdown and cost comparison',
     'Fig. 17.  BOM breakdown and cost comparison'),
    ('Fig. 17.  QR  Code for Github Repo',
     'Fig. 18.  QR  Code for Github Repo'),
]
n = 0
for p in paras:
    t = ptext(p)
    nt = t
    for old, new in ref_fixes:
        if old in nt:
            nt = nt.replace(old, new)
    if nt != t:
        set_text(p, nt)
        n += 1
print('renumber refs applied to', n, 'paragraphs')

# ── 3) limitations sentence ────────────────────────────────────────────────
lim = next(p for p in paras if ptext(p).strip().startswith(
    'Limitations of this study include'))
lim_text = ptext(lim)
assert lim_text.endswith('identical simulator conditions.')
set_text(lim, lim_text + ' In addition, the execution-time measurement of '
         'Section V covers the Helios side only; the STM32F103C8T6 receive, '
         'parse and PWM-update latency and the physical end-to-end latency '
         'across the UART link remain to be instrumented.')
print('limitations sentence appended')

# ── 4) future work sentence ────────────────────────────────────────────────
fw = next(p for p in paras if ptext(p).strip().startswith(
    'Several directions for future work arise'))
fw_text = ptext(fw)
set_text(fw, fw_text + ' Fifth, the closed-loop latency chain from irradiance '
         'sample to PWM edge will be measured end to end with the Artemis '
         '(STM32F103C8T6) board instrumented, closing the remaining gap in '
         'Table V.')
print('future work sentence appended')

doc.save(DOC)
print('revision23 saved')