"""Convert response_letter.md to a clean DOCX (Times New Roman, IJPEDS-style)."""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SRC = '/home/touhid/artemis-helios/response_letter.md'
OUT = '/home/touhid/artemis-helios/Docs/Upload_package/response_letter.docx'


def set_font(run, size=11, bold=False, italic=False, mono=False, color=None):
    run.font.name = 'Liberation Serif' if mono else 'Times New Roman'
    r = run._element.rPr
    if r is None:
        run._element.get_or_add_rPr()
        r = run._element.rPr
    rf = r.find(qn('w:rFonts'))
    if rf is None:
        rf = r.makeelement(qn('w:rFonts'), {})
        r.insert(0, rf)
    rf.set(qn('w:ascii'), 'Liberation Serif' if mono else 'Times New Roman')
    rf.set(qn('w:hAnsi'), 'Liberation Serif' if mono else 'Times New Roman')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_runs(p, text, size=11):
    """Parse **bold**, *italic*, `code` inline markers into runs."""
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            set_font(p.add_run(tok[2:-2]), size=size, bold=True)
        elif tok.startswith('`') and tok.endswith('`'):
            set_font(p.add_run(tok[1:-1]), size=size, mono=True)
        elif tok.startswith('*') and tok.endswith('*') and len(tok) > 2:
            set_font(p.add_run(tok[1:-1]), size=size, italic=True)
        else:
            set_font(p.add_run(tok), size=size)


def heading(doc, text, size):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if size >= 13 else 10)
    p.paragraph_format.space_after = Pt(6)
    add_runs(p, text, size=size)
    for r in p.runs:
        r.font.bold = True
    return p


def body(doc, text, size=11, indent=0.0, italic=False, bullet=None, num=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    add_runs(p, text, size=size)
    if italic:
        for r in p.runs:
            r.font.italic = True
    return p


doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

with open(SRC, encoding='utf-8') as f:
    lines = f.read().split('\n')

i = 0
in_table = False
table_rows = []
n = len(lines)
while i < n:
    line = lines[i].rstrip()
    if not line.strip():
        i += 1
        continue
    # table
    if line.lstrip().startswith('|'):
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        sep = all(re.fullmatch(r':?-{2,}:?', c.replace(' ', '')) for c in cells)
        if not sep:
            table_rows.append(cells)
        i += 1
        continue
    if line.lstrip().startswith('# ') :
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        add_runs(p, line.lstrip()[2:].strip(), size=16)
        for r in p.runs: r.font.bold = True
    elif line.lstrip().startswith('## '):
        heading(doc, line.lstrip()[3:].strip(), size=14)
    elif line.lstrip().startswith('### '):
        heading(doc, line.lstrip()[4:].strip(), size=12)
    elif line.strip() == '---':
        pass  # horizontal rule: blank line separator
    elif line.lstrip().startswith('> '):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.space_after = Pt(6)
        add_runs(p, line.lstrip()[2:].strip(), size=10.5)
        for r in p.runs: r.font.italic = True
    elif re.match(r'^\s*\d+\.\s+', line):
        m = re.match(r'^\s*(\d+)\.\s+(.*)$', line)
        txt = m.group(2)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(f'{m.group(1)}.  ')
        set_font(run, size=11, bold=True)
        add_runs(p, txt, size=11)
    elif line.lstrip().startswith('- '):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run('•  ')
        set_font(run, size=11)
        add_runs(p, line.lstrip()[2:].strip(), size=11)
    elif line.lstrip().startswith('**'):
        body(doc, line.strip(), size=11)
    else:
        body(doc, line.strip(), size=11)
    i += 1

# build the summary table if rows collected
if table_rows:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    header = doc.add_paragraph()
    run = header.add_run('Summary of Changes')
    set_font(run, size=12, bold=True)
    t = doc.add_table(rows=len(table_rows), cols=max(len(r) for r in table_rows))
    t.style = 'Table Grid'
    for ri, row in enumerate(table_rows):
        for ci, cell in enumerate(row):
            c = t.cell(ri, ci)
            c.paragraphs[0].paragraph_format.space_after = Pt(2)
            add_runs(c.paragraphs[0], cell, size=9.5)
            if ri == 0:
                for r in c.paragraphs[0].runs:
                    r.font.bold = True

doc.save(OUT)
print('saved', OUT)