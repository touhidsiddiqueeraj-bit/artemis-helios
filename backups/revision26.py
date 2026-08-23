"""
revision26.py — Artemis STM32 timing (Table V expansion + Fig.18 dual-MCU)
==========================================================================
1) Expand Table V (docx table 4) from Helios-only (8 rows) to dual-MCU:
   keep Helios block, add Artemis block (INA219, parse, VS-P&O+blend+PWM,
   UART TX, full tick, loop period, end-to-end) with DWT_CYCCNT stats
   from Code/Python/results/artemis_timing.csv (N=400, 72 MHz).
2) Replace Fig.18 image (fig_timing_budget.png) with new dual-MCU version.
3) Update Fig.18 caption and Section V paragraph to mention both sides
   and 3.55 ms end-to-end latency.
Run: python3 backups/revision26.py
"""
import re, copy, os
from docx import Document
from docx.shared import Emu
from docx.text.paragraph import Paragraph
from docx.oxml import parse_xml

DOC='25195-52952-1-SM-REVISED.docx'
W_='{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

doc=Document(DOC)

# ── 1) Table V expansion ────────────────────────────────────────────
# docx table 4 is Table V
tbl=doc.tables[4]
# Verify header
assert tbl.cell(0,0).text.strip().startswith('Stage')
# New rows to insert after current 8 rows (after jitter row)
new_rows=[
    # Stage, Mean, p99, Max  (all strings, µs/ms)
    ('INA219 read (8-sample, 400 kHz)', '8.517 ms', '8.723 ms', '8.785 ms'),
    ('UART parse (HEL frame)', '22.5 µs', '31.6 µs', '37.1 µs'),
    ('VS-P&O + blend + PWM update', '24.9 µs', '34.2 µs', '38.0 µs'),
    ('UART TX Artemis→Helios (30 B)', '2.610 ms', '2.647 ms', '2.654 ms'),
    ('Full Artemis control tick', '11.256 ms', '11.456 ms', '11.539 ms'),
    ('Loop period Artemis (100 ms nom.)', '99.999 ms', '100.058 ms', '100.058 ms'),
    ('End-to-end Helios UART→Artemis PWM', '3.55 ms', '—', '—'),
]
# Add a separator row label? Instead just append rows
for stage, mean, p99, mx in new_rows:
    row=tbl.add_row()
    # keep same style: 4 cells
    row.cells[0].text=stage
    row.cells[1].text=mean
    row.cells[2].text=p99
    row.cells[3].text=mx
    # small font: set cell text size 8pt via run props (optional, keep default for now)
    for cell in row.cells:
        for p in cell.paragraphs:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size=None

# Update Table V caption paragraph (preceding the table? Actually caption is a paragraph before table)
# Find paragraph with "Table V. Measured Helios"
cap=None
for p in doc.paragraphs:
    if p.text.strip().startswith('Table V. Measured Helios'):
        cap=p
        break
if cap:
    cap.text='Table V. Measured dual-MCU execution-time budget (Helios ESP32-S3 N16R8 240 MHz and Artemis STM32F103C8T6 72 MHz, N = 400 each; Helios via esp_timer, Artemis via DWT_CYCCNT through ESP32 UART bridge). End-to-end Helios UART TX → Artemis parse → PWM update = 3.55 ms.'
    print('Table V caption updated')
print(f'Table V expanded: now {len(tbl.rows)} rows')

# ── 2) Fig.18 caption update ───────────────────────────────────────
for p in doc.paragraphs:
    if p.text.strip().startswith('Fig. 18. Measured Helios execution-time'):
        p.text='Fig. 18. Measured dual-MCU execution-time budget (N = 400 each): (a) Helios (ESP32-S3, 240 MHz) and Artemis (STM32F103C8T6, 72 MHz, DWT_CYCCNT) control ticks against the 100 ms budget (Artemis 11.26 ms, Helios 10.00 ms, 88–90 ms idle); (b) Artemis 100 ms loop-period distribution (mean 100.00 ms, p99 100.06 ms, max 100.06 ms). The 3.55 ms Helios-UART→Artemis-PWM path is an order of magnitude below the 5 s cloud-edge transient.'
        print('Fig.18 caption updated')
        break

# ── 3) Section V paragraph patch — add Artemis sentence after Helios timing paragraph ──
# Find paragraph that contains "LSTM inference (24 × 32 units) 6.355 ms" is in table, not paragraph.
# Section V text is around "Measured Helios execution-time budget" narrative.
# We patch the paragraph that starts with "The timing budget confirms"
target=None
for p in doc.paragraphs:
    if 'The timing budget confirms' in p.text or 'timing budget on the ESP32' in p.text:
        target=p
        break
if target:
    # Append Artemis sentence
    extra=' On the Artemis side (STM32F103C8T6, 72 MHz, DWT_CYCCNT, N = 400, 8-sample INA219 @400 kHz), the 100 ms tick averages 11.26 ms (INA219 8.52 ms, UART parse 22.5 µs, VS-P&O+blend+PWM 24.9 µs, UART TX 2.61 ms; p99 11.46 ms, max 11.54 ms) with loop jitter p99 58 µs; the Helios UART TX (3.48 ms) → Artemis parse (22.5 µs) → PWM update (0.8 µs) end-to-end latency is 3.55 ms, well within the 5 s transient window and the 100 ms control deadline.'
    if 'Artemis side' not in target.text:
        target.text = target.text.rstrip() + extra
        print('Section V paragraph patched')

doc.save(DOC)
print('revision26 saved')
