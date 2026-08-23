"""
revision24.py — controlled transient benchmark suite (Table VI, new Fig. 15)
==============================================================================
1) Renumber figures 15..18 -> 16..19 (field logger, timing budget, BOM, QR).
2) Insert at the end of Section IV (before "V. HARDWARE IMPLEMENTATION"):
   - subsection heading "J.  Controlled Transient Benchmark Suite"
   - intro paragraph (suite definition, fairness, conservative-forecast note)
   - Table VI (6 waveforms x 4 controllers x 8 metrics, from
     Code/Python/results/transient_benchmark.csv)
   - Fig. 15 (Figures/fig15_transient_benchmark.png) with caption

Run:  python3 backups/revision24.py
"""
import re
import copy
import csv
import os
from lxml import etree
from docx import Document
from docx.shared import Emu
from docx.text.paragraph import Paragraph
from docx.oxml import parse_xml

DOC = '25195-52952-1-SM-REVISED.docx'
W_ = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
COL_W_EMU = 3187700
COL_W_TWIPS = 5220          # single column: 5230 twips usable
FIG = 'Figures/fig15_transient_benchmark.png'
CSV = 'Code/Python/results/transient_benchmark.csv'

doc = Document(DOC)
body = doc.element.body
paras = list(doc.paragraphs)


def ptext(p):
    return ''.join(r.text or '' for r in p.runs)


def set_text(p, new):
    for r in p.runs:
        r.text = ''
    p.runs[0].text = new


def replace_run_text(el, text):
    for r in el.findall(f'{W_}r'):
        el.remove(r)
    run = etree.SubElement(el, f'{W_}r')
    rPr = etree.SubElement(run, f'{W_}rPr')
    sz = etree.SubElement(rPr, f'{W_}sz'); sz.set(W_ + 'val', '20')
    szcs = etree.SubElement(rPr, f'{W_}szCs'); szcs.set(W_ + 'val', '20')
    t = etree.SubElement(run, f'{W_}t'); t.set(W_ + 'space', 'preserve')
    t.text = text
    return el


def image_para():
    ns = W_.strip('{}')
    xml = (f'<w:p xmlns:w="{ns}">'
           f'<w:pPr><w:jc w:val="center"/></w:pPr>'
           f'<w:r/></w:p>')
    return parse_xml(xml)


# ── 1) renumber Fig. 15..18 -> 16..19 (captions and in-text refs) ───────────
renum = re.compile(r'Fig\. (15|16|17|18)')
n = 0
for p in paras:
    t = ptext(p)
    nt = renum.sub(lambda m: 'Fig. ' + str(int(m.group(1)) + 1), t)
    if nt != t:
        set_text(p, nt)
        n += 1
print('renumbered Fig. refs in', n, 'paragraphs')

# ── 2) insertion before the Section V heading ──────────────────────────────
v_head = next(p for p in paras if ptext(p).strip().startswith(
    'V. HARDWARE IMPLEMENTATION'))
anchor = v_head._p

# subsection heading, cloned from "I.  Sensitivity Analysis"
heading_el = next(p for p in paras if ptext(p).strip().startswith(
    'I.  Sensitivity Analysis'))._p
head = copy.deepcopy(heading_el)
replace_run_text(head, 'J.  Controlled Transient Benchmark Suite')

# body paragraph, cloned from the Table V caption paragraph (Normal, 10 pt)
tmpl = next(p for p in paras if ptext(p).strip().startswith(
    'Table V. Measured Helios execution-time budget'))

INTRO = ('A controlled transient suite completes the controller comparison '
         'under exactly the identical conditions of Section IV.B: the same '
         'single-diode PV model, 0.1 s sampling, sensing resolution and '
         'initialization for every controller, with no per-controller '
         're-tuning. Six irradiance waveforms are exercised \u2014 step-up '
         '(600\u21921000 W/m\u00b2), step-down (1000\u2192600 W/m\u00b2), a 30 s '
         'ramp (400\u2192900 W/m\u00b2), a cloud-edge drop (900\u2192300 W/m\u00b2 '
         'over 5 s with a 20 s recovery), five repeated cloud edges at 60 s '
         'spacing, and the stochastic Markov+OU day (seed 23). For each '
         'waveform, Table VI and Fig. 15 report tracking efficiency, maximum '
         'tracking error, 2%-band settling time, overshoot and undershoot, '
         'energy not captured, steady-state oscillation amplitude and mean '
         'MPP voltage error. The LSTM variant is fed a causal low-pass '
         'forecast of the measured irradiance rather than the trained '
         'predictor, so its predictive branch carries no skill on '
         'sub-minute transients; the suite is therefore a conservative lower '
         'bound on the LSTM-assisted controller. Even under this handicap '
         'Helios-Artemis meets or exceeds fixed-step P&O on every waveform '
         '\u2014 95.1% versus 89.5% on the stochastic day and 87.5% versus '
         '87.4% on the cloud edge \u2014 with both step transients settling '
         'within 1.8 s. The incremental-conductance variant recovers fastest '
         'on this synthetic stress day (96.9%), which exercises a smooth P\u2013V '
         'surface without the diurnal and battery constraints of the '
         'calibrated full-day profiles; those profiles remain the scenario '
         'of Table III.')

intro = copy.deepcopy(tmpl._p)
replace_run_text(intro, INTRO)
tcap = copy.deepcopy(tmpl._p)
replace_run_text(tcap, 'Table VI. Controlled transient benchmark results '
                       '(dt = 0.1 s, seed 23; units: tracking '
                       'efficiency, error, overshoot and undershoot '
                       'in %, settling time in s, energy in mWh, '
                       'oscillation in W, voltage error in V. '
                       'P&O = fixed-step P&O, VS = variable-step P&O, '
                       'INC = incremental conductance, LSTM = '
                       'LSTM-assisted P&O).')
pPr = tcap.find(f'{W_}pPr')
pb = etree.Element(f'{W_}pageBreakBefore')
pPr.insert(0, pb)   # start caption+table at the top of a fresh page
fcap = copy.deepcopy(tmpl._p)
replace_run_text(fcap, 'Fig. 15. Controlled transient benchmark across all '
                       'four controllers under identical conditions: '
                       '(a) post-step power traces on the step-down waveform '
                       'against the analytic P\u2009MPP, (b) tracking '
                       'efficiency per waveform (pattern-coded). Values are '
                       'summarised in Table VI.')


def fmt_metric(name, v):
    if name == 'eta_track':
        return f'{v:.1f}'
    if name in ('max_tracking_error', 'overshoot', 'undershoot',
                'mpp_voltage_error'):
        return str(int(v)) if v == int(v) and abs(v) >= 100 else f'{v:.1f}'
    if name == 'settling_time':
        return 'N/A' if v != v else f'{v:.1f}'
    if name == 'energy_not_captured':
        v = v * 1000.0  # Wh -> mWh
        return f'{v:.0f}' if v >= 10 else f'{v:.2f}'
    return f'{v:.3f}'


def build_table():
    with open(CSV) as f:
        rows = list(csv.DictReader(f))
    assert len(rows) == 6 * 4 * 8, 'CSV shape mismatch'
    waves = ['step-up', 'step-down', 'ramp', 'cloud-edge',
             'repeated-cloud', 'stochastic']
    ctrls = ['Plain-P&O', 'VS-P&O', 'INC', 'LSTM-P&O']
    disp = {'Plain-P&O': 'Plain-P&O', 'VS-P&O': 'VS-P&O', 'INC': 'INC', 'LSTM-P&O': 'LSTM'}
    metrics = ['eta_track', 'max_tracking_error', 'settling_time',
               'overshoot', 'undershoot', 'energy_not_captured',
               'oscillation_amplitude', 'mpp_voltage_error']
    val = {(r['waveform'], r['controller'], r['metric']): float(r['value'])
           for r in rows}
    headers = ['\u03b7', 'e_max', 't_s', 'over', 'under',
               'E', 'osc', 'V_err']

    tbl = etree.Element(f'{W_}tbl')
    tblPr = etree.SubElement(tbl, f'{W_}tblPr')
    tblW = etree.SubElement(tblPr, f'{W_}tblW')
    tblW.set(W_ + 'w', str(COL_W_TWIPS)); tblW.set(W_ + 'type', 'dxa')
    etree.SubElement(tblPr, f'{W_}tblLayout').set(W_ + 'type', 'fixed')
    etree.SubElement(tblPr, f'{W_}jc').set(W_ + 'val', 'center')
    borders = etree.SubElement(tblPr, f'{W_}tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = etree.SubElement(borders, f'{W_}{side}')
        b.set(W_ + 'val', 'single'); b.set(W_ + 'sz', '4')
        b.set(W_ + 'color', '000000')
    cellmar = etree.SubElement(tblPr, f'{W_}tblCellMar')
    for m in ('left', 'start', 'right', 'end'):
        cm = etree.SubElement(cellmar, f'{W_}{m}')
        cm.set(W_ + 'w', '30'); cm.set(W_ + 'type', 'dxa')
    widths = [1060, 640] + [440] * 8
    assert sum(widths) == COL_W_TWIPS
    grid = etree.SubElement(tbl, f'{W_}tblGrid')
    for wd in widths:
        etree.SubElement(grid, f'{W_}gridCol').set(W_ + 'w', str(wd))

    def row(cells):
        tr = etree.SubElement(tbl, f'{W_}tr')
        trPr = etree.SubElement(tr, f'{W_}trPr')
        etree.SubElement(trPr, f'{W_}cantSplit')
        for ci, v in enumerate(cells):
            tc = etree.SubElement(tr, f'{W_}tc')
            tcPr = etree.SubElement(tc, f'{W_}tcPr')
            tcW = etree.SubElement(tcPr, f'{W_}tcW')
            tcW.set(W_ + 'w', str(widths[ci])); tcW.set(W_ + 'type', 'dxa')
            etree.SubElement(tcPr, f'{W_}vAlign').set(W_ + 'val', 'center')
            p = etree.SubElement(tc, f'{W_}p')
            r = etree.SubElement(p, f'{W_}r')
            rPr = etree.SubElement(r, f'{W_}rPr')
            sz = etree.SubElement(rPr, f'{W_}sz'); sz.set(W_ + 'val', '12' if ci == 1 else '13')
            t = etree.SubElement(r, f'{W_}t')
            t.set(W_ + 'space', 'preserve')
            t.text = v

    row(['Waveform', 'Ctrl'] + headers)
    for w in waves:
        for c in ctrls:
            row([w, disp[c]] + [fmt_metric(m, val[(w, c, m)]) for m in metrics])
    return tbl


for el in (head, intro, tcap, build_table()):
    anchor.addprevious(el)
fig_p = image_para()
Paragraph(fig_p, doc).add_run().add_picture(FIG, width=Emu(COL_W_EMU))
anchor.addprevious(fig_p)
anchor.addprevious(fcap)
print('IV.J block inserted (heading, paragraph, Table VI, Fig. 15)')

doc.save(DOC)
print('revision24 saved')
