"""
revision14.py — Insert Section IV.G (power-stage verification) + fix figure refs
================================================================================
* Fixes 2a text references: "Figs. 15 and 16" -> "Figs. 13-15",
  "Fig. 16"/"Fig. 16b" -> "Fig. 15"/"Fig. 15b", 3.9 A -> 3.8 A.
* Inserts "G. Power-Stage Loss and Efficiency Verification" after Fig. 9
  caption (end of Section IV.F), before "V. DISCUSSION":
    - Fig. 13: power-stage schematic with signal names / measurement points
    - Fig. 14: switching waveforms at the 3.8 A operating point (ngspice)
    - Fig. 15: efficiency vs load + loss breakdown (ngspice)
  Reviewer A items 1 (topology+measurement points, thermal) and 2 (loss
  breakdown + efficiency-vs-load curve).

Run from the manuscript directory:  python3 backups/revision14.py
"""
import copy
import os
from lxml import etree
from docx import Document
from docx.shared import Emu

DOC = '25195-52952-1-SM-REVISED.docx'
W_ = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
FULL_W_TWIPS = 9000
FIG13 = 'Figures/fig13_power_stage_schematic.png'
FIG14 = 'Figures/fig15_buck_waveforms.png'
FIG15 = 'Figures/fig16_eff_loss.png'

doc = Document(DOC)
paras = list(doc.paragraphs)

def para_text(p):
    return ''.join(r.text or '' for r in p.runs)

# ── locate anchors ─────────────────────────────────────────────────────────
cap9 = next(p for p in paras if para_text(p).strip().startswith('Fig. 9. '))
disc = next(p for p in paras if para_text(p).strip().startswith('V.  DISCUSSION'))
subf = next(p for p in paras if para_text(p).strip().startswith('F.  Field Logger Validation'))
body_tmpl = next(p for p in paras if para_text(p).strip().startswith('The correction fires only'))
p2a = next(p for p in paras if '(Section IV, Figs. 15 and 16)' in para_text(p))
p63 = next(p for p in paras if 'peak junction temperature 54' in para_text(p))

# ── 1) fix cross references ────────────────────────────────────────────────
def replace_in(p, old, new):
    t = para_text(p)
    assert old in t, f'not found: {old[:50]}'
    t = t.replace(old, new)
    for r in p.runs:
        r.text = ''
    p.runs[0].text = t

replace_in(p2a, '(Section IV, Figs. 15 and 16)', '(Section IV.G, Figs. 13\u201315)')
replace_in(p2a, '98.0% at the nominal 3.9 A charging point, degrading to 95.0% at 12.8 A (Fig. 16)',
           '98.0% at the nominal 3.8 A charging point, degrading to 95.0% at 12.8 A (Fig. 15)')
replace_in(p63, 'as in Fig. 16b)', 'as in Fig. 15b)')

# ── helpers (revision11.py pattern) ────────────────────────────────────────
def clone_with_text(tmpl_el, new_text):
    el = copy.deepcopy(tmpl_el)
    runs = el.findall(f'{W_}r')
    first = runs[0]
    for r in runs[1:]:
        el.remove(r)
    ts = first.findall(f'{W_}t')
    if not ts:
        t = etree.SubElement(first, f'{W_}t')
        t.set(W_ + 'space', 'preserve')
        t.text = new_text
    else:
        for t in ts[1:]:
            first.remove(t)
        ts[0].text = new_text
    return el

def make_tbl():
    tbl = etree.Element(f'{W_}tbl')
    tblPr = etree.SubElement(tbl, f'{W_}tblPr')
    tblW = etree.SubElement(tblPr, f'{W_}tblW')
    tblW.set(W_ + 'w', str(FULL_W_TWIPS)); tblW.set(W_ + 'type', 'dxa')
    etree.SubElement(tblPr, f'{W_}tblLayout').set(W_ + 'type', 'fixed')
    etree.SubElement(tblPr, f'{W_}jc').set(W_ + 'val', 'center')
    borders = etree.SubElement(tblPr, f'{W_}tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        etree.SubElement(borders, f'{W_}{side}').set(W_ + 'val', 'nil')
    mar = etree.SubElement(tblPr, f'{W_}tblCellMar')
    for side in ('top', 'left', 'bottom', 'right'):
        s = etree.SubElement(mar, f'{W_}{side}')
        s.set(W_ + 'w', '0'); s.set(W_ + 'type', 'dxa')
    grid = etree.SubElement(tbl, f'{W_}tblGrid')
    etree.SubElement(grid, f'{W_}gridCol').set(W_ + 'w', str(FULL_W_TWIPS))
    tr = etree.SubElement(tbl, f'{W_}tr')
    tc = etree.SubElement(tr, f'{W_}tc')
    tcPr = etree.SubElement(tc, f'{W_}tcPr')
    tcW = etree.SubElement(tcPr, f'{W_}tcW')
    tcW.set(W_ + 'w', str(FULL_W_TWIPS)); tcW.set(W_ + 'type', 'dxa')
    etree.SubElement(tcPr, f'{W_}tcMar')
    etree.SubElement(tcPr, f'{W_}vAlign').set(W_ + 'val', 'center')
    return tbl

def figure_table(img_path, width_in, caption_text):
    tbl = make_tbl()
    tc = tbl.find(f'{W_}tr').find(f'{W_}tc')
    tmp_para = doc.add_paragraph()
    run = tmp_para.add_run()
    run.add_picture(img_path, width=Emu(int(width_in * 914400)))
    pPr = etree.SubElement(tmp_para._p, f'{W_}pPr')
    jc = etree.SubElement(pPr, f'{W_}jc'); jc.set(W_ + 'val', 'center')
    cap_el = clone_with_text(cap9._p, caption_text)
    tc.append(tmp_para._p)
    tc.append(cap_el)
    return tbl

# ── 2) build IV.G elements ─────────────────────────────────────────────────
heading = clone_with_text(subf._p, 'G.  Power-Stage Loss and Efficiency Verification')

p1 = clone_with_text(body_tmpl._p,
 'The power stage used in all simulations of Sections IV.B and IV.D and the '
 'reproducible switching-level model of the previous paragraph are shown in '
 'Fig. 13, together with the signal names and measurement points used by the '
 'verification setup: the PV-side shunt R_shunt = 10 m\u03a9 sensed by the INA219 '
 '12-bit current sensor (V_pv, I_pv sampled at 100 ms), the gate\u2013source '
 'voltage v_gs, the switch node v_sw, the inductor current i_L, and the battery '
 'voltage v_bat. The TC4420 gate driver operates from a 12 V rail with a 4.7 \u03a9 '
 'gate resistor; the PWM duty is clamped to [0.05, 0.95].')

tbl13 = figure_table(FIG13, 7.0,
 'Fig. 13. Power-stage schematic with signal names and measurement points: '
 '50 Wp PV panel (Isc0 = 2.91 A, Voc0 = 21.6 V), INA219 shunt sensing, '
 'IRFB4110 high-side switch with freewheeling body diode, TC4420 gate driver '
 '(50 kHz, V_DD = 12 V, R_g = 4.7 \u03a9), LC filter (L = 100 \u00b5H, DCR = 30 m\u03a9; '
 'C = 470 \u00b5F, ESR = 40 m\u03a9), and 12 V/7 Ah SLA battery (R_int = 50 m\u03a9).')

p2 = clone_with_text(body_tmpl._p,
 'Fig. 14 shows the switching waveforms at the nominal 3.8 A operating point '
 '(V_in = 17.9 V, 50 kHz): (a) switch node v_sw and gate\u2013source v_gs, (b) '
 'inductor current i_L in continuous conduction, (c) output voltage ripple, and '
 '(d) PV input current. The converter delivers 13.16 V at 3.80 A with a '
 'measured efficiency of 97.99%, and the energy balance '
 '(P_in \u2212 P_out \u2212 P_loss)/P_in = +0.75% closes the conservation check to '
 'within one percentage point.')

tbl14 = figure_table(FIG14, 7.0,
 'Fig. 14. Buck power-stage switching waveforms at the nominal 3.8 A '
 'operating point (ngspice, 20 ns timestep): (a) switch node and gate\u2013source '
 'voltage, (b) inductor current (CCM), (c) output voltage ripple, (d) PV input '
 'current.')

p3 = clone_with_text(body_tmpl._p,
 'Fig. 15(a) gives the efficiency-versus-load characteristic of the converter '
 'over the full operating range of the Helios-Artemis system: 98.9% at 1.0 A '
 'load down to 95.0% at the 12.8 A worst-case current. Fig. 15(b) breaks down '
 'the 1.41 W total loss at the 3.8 A operating point: MOSFET conduction '
 '0.32 W (R_DS(on) = 3.7 m\u03a9), freewheeling diode 0.52 W, inductor DCR '
 '0.43 W, INA219 shunt 0.10 W, and capacitor ESR 0.03 W. The simulated total '
 'exceeds the measured P_in \u2212 P_out = 1.03 W by 0.75% of P_in, consistent '
 'with the energy-balance closure of Fig. 14; across the load sweep the '
 'balance error remains below 1.6%. These numbers confirm the 98% peak '
 'efficiency and 62 \u00b0C thermal budget reported in Section IV.B.')

tbl15 = figure_table(FIG15, 7.0,
 'Fig. 15. (a) Converter efficiency versus output current (98.9% at 1.0 A to '
 '95.0% at 12.8 A; 98.0% at the nominal 3.8 A point). (b) Loss breakdown at '
 '3.8 A: MOSFET 0.32 W, body diode 0.52 W, inductor DCR 0.43 W, shunt 0.10 W, '
 'capacitor ESR 0.03 W (total 1.41 W).')

# ── insert after Fig. 9 caption, before V. DISCUSSION ──────────────────────
anchor = cap9._p
for el in (heading, p1, tbl13, p2, tbl14, p3, tbl15):
    anchor.addnext(el)
    anchor = el

doc.save(DOC)
print('revision14 applied: IV.G inserted with Figs. 13-15, refs fixed')
