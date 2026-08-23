"""
letter_to_docx.py — Convert response_letter_v2.md to .docx
===========================================================
Targeted converter for the round-2 response letter (known structure):
# / ## / ### headings, > blockquote reviewer quotes, **bold**, lists, pipe tables.
Output: 25195-52952-1-SM-REVISED-RESPONSE-LETTER.docx
"""
import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

SRC = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'response_letter_v2.md')
OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..',
                   '25195-52952-1-SM-REVISED-RESPONSE-LETTER.docx')

doc = Document()
# base font
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

def add_runs(p, text):
    """Handle **bold** and *italic* inline."""
    pos = 0
    for m in re.finditer(r'\*\*(.+?)\*\*|\*(.+?)\*', text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        r = p.add_run(m.group(1) or m.group(2))
        r.bold = bool(m.group(1))
        r.italic = bool(m.group(2))
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])

def para(text, style=None):
    p = doc.add_paragraph()
    if style:
        p.style = doc.styles[style]
    add_runs(p, text)
    return p

lines = open(SRC).read().splitlines()
i = 0
i_hdr = 0
while i < len(lines):
    line = lines[i].rstrip()
    if not line.strip():
        i += 1
        continue
    if line.startswith('|'):
        # table: consume consecutive pipe lines
        tbl = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            tbl.append(lines[i].strip()); i += 1
        rows = [[c.strip() for c in r.strip('|').split('|')] for r in tbl]
        rows = [r for r in rows if not (all(c.replace('-','').replace(':','').strip()=='' for c in r))]
        if rows:
            t = doc.add_table(rows=len(rows), cols=len(rows[0]))
            t.style = 'Light Grid Accent 1'
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    cp = t.cell(ri, ci).paragraphs[0]
                    add_runs(cp, cell)
                    if ri == 0:
                        cp.runs[0].bold = True
        continue
    if line.startswith('### '):
        para(line[4:], 'Heading 3'); i += 1; continue
    if line.startswith('## '):
        para(line[3:], 'Heading 2'); i += 1; continue
    if line.startswith('# '):
        para(line[2:], 'Title'); i += 1; continue
    if line.strip() == '---':
        i += 1; continue
    if line.startswith('> '):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.right_indent = Inches(0.2)
        add_runs(p, line[2:])
        for r in p.runs:
            r.italic = True
            r.font.size = Pt(10)
        # make reviewer-quote prefix bold marker mild
        i += 1
        continue
    if re.match(r'^\d+\.\s', line):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        add_runs(p, line)
        i += 1
        continue
    if line.startswith('- '):
        p = doc.add_paragraph(style='List Bullet')
        add_runs(p, line[2:])
        i += 1
        continue
    # plain paragraph
    para(line)
    i += 1

doc.save(OUT)
print('saved', OUT, os.path.getsize(OUT), 'bytes')
