"""
revision15.py — Fix figure embedding for the two-column template
================================================================
Round-1 figures are inline single-column images (3.486 in, centered
paragraph) with standalone caption paragraphs. My earlier insertion
(revision14.py) used full-width tables (9000 twips) with 7.0-in images,
which overflow the single column and overlap the neighbouring column
(and the graphical-abstract table did the same). This script:
  * removes the 4 tables (Graphical Abstract + Figs 13-15),
  * re-embeds the images inline at column width (3187700 EMU),
  * appends standalone caption paragraphs (round-1 pattern),
  * shortens the captions.

Run:  python3 backups/revision15.py
"""
import copy
from lxml import etree
from docx import Document
from docx.shared import Emu

DOC = '25195-52952-1-SM-REVISED.docx'
W_ = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
COL_W_EMU = 3187700  # 3.486 in — round-1 figure width
GA = 'Figures/graphical_abstract.png'
FIG13 = 'Figures/fig13_power_stage_schematic.png'
FIG14 = 'Figures/fig15_buck_waveforms.png'
FIG15 = 'Figures/fig16_eff_loss.png'

CAP_GA = 'Graphical Abstract'
CAP13 = ('Fig. 13. Power-stage schematic with signal names and measurement points '
         '(V_pv, I_pv via INA219 shunt, v_gs, v_sw, i_L, v_bat): 50 Wp PV panel, '
         'IRFB4110 with freewheeling body diode, TC4420 gate driver '
         '(50 kHz, 12 V, R_g = 4.7 \u03a9), LC filter (L = 100 \u00b5H, DCR 30 m\u03a9; '
         'C = 470 \u00b5F, ESR 40 m\u03a9), 12 V/7 Ah SLA battery (R_int = 50 m\u03a9).')
CAP14 = ('Fig. 14. Buck power-stage switching waveforms at the nominal 3.8 A '
         'operating point (ngspice, 20 ns timestep): (a) switch node and '
         'gate\u2013source voltage, (b) inductor current (CCM), (c) output voltage '
         'ripple, (d) PV input current.')
CAP15 = ('Fig. 15. (a) Converter efficiency versus output current (98.9% at 1.0 A '
         'to 95.0% at 12.8 A; 98.0% at the 3.8 A point). (b) Loss breakdown at '
         '3.8 A: MOSFET 0.32 W, body diode 0.52 W, inductor DCR 0.43 W, shunt '
         '0.10 W, capacitor ESR 0.03 W (total 1.41 W).')

doc = Document(DOC)
body = doc.element.body

def para_text(p):
    return ''.join(r.text or '' for r in p.runs)

# ── locate the four tables by their caption text ───────────────────────────
def table_caption(tbl):
    return ''.join(t.text or '' for t in tbl.iter() if t.tag.endswith('}t')).strip()

to_remove = []
for tbl in list(body.iter(f'{W_}tbl')):
    cap = table_caption(tbl)
    if cap.startswith('Graphical Abstract') or cap.startswith('Fig. 13.') \
       or cap.startswith('Fig. 14.') or cap.startswith('Fig. 15.'):
        to_remove.append(tbl)
print('removing tables:', len(to_remove))
for tbl in to_remove:
    body.remove(tbl)

# ── anchors ────────────────────────────────────────────────────────────────
paras = list(doc.paragraphs)
cap9 = next(p for p in paras if para_text(p).strip().startswith('Fig. 9. '))
intro = next(p for p in paras if para_text(p).strip().startswith('I.  INTRODUCTION'))
p1 = next(p for p in paras if para_text(p).strip().startswith('The power stage used in all simulations'))
p2 = next(p for p in paras if para_text(p).strip().startswith('Fig. 14 shows the switching waveforms'))
p3 = next(p for p in paras if para_text(p).strip().startswith('Fig. 15(a) gives the efficiency-versus-load'))

# ── builders (round-1 pattern) ─────────────────────────────────────────────
def caption_el(text):
    el = copy.deepcopy(cap9._p)
    runs = el.findall(f'{W_}r')
    first = runs[0]
    for r in runs[1:]:
        el.remove(r)
    ts = first.findall(f'{W_}t')
    if not ts:
        t = etree.SubElement(first, f'{W_}t')
        t.set(W_ + 'space', 'preserve')
        t.text = text
    else:
        for t in ts[1:]:
            first.remove(t)
        ts[0].text = text
    return el

def image_el(path):
    tmp = doc.add_paragraph()
    pPr = etree.SubElement(tmp._p, f'{W_}pPr')
    jc = etree.SubElement(pPr, f'{W_}jc'); jc.set(W_ + 'val', 'center')
    tmp.add_run().add_picture(path, width=Emu(COL_W_EMU))
    return tmp._p

# ── insert: GA after Keywords (before Introduction) ────────────────────────
intro._p.addprevious(image_el(GA))
intro._p.addprevious(caption_el(CAP_GA))

# ── insert: IV.G figures ───────────────────────────────────────────────────
p1._p.addnext(caption_el(CAP13))
p1._p.addnext(image_el(FIG13))
p2._p.addnext(caption_el(CAP14))
p2._p.addnext(image_el(FIG14))
p3._p.addnext(caption_el(CAP15))
p3._p.addnext(image_el(FIG15))

doc.save(DOC)
print('revision15 applied: 4 tables removed, images re-embedded inline at column width')
