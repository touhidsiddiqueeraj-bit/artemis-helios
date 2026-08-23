from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

DOC = '25195-52952-1-SM-REVISED.docx'


def link(paragraph, label, url):
    rid = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('r:id'), rid)
    run = OxmlElement('w:r')
    props = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    props.append(color)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    props.append(underline)
    run.append(props)
    text = OxmlElement('w:t')
    text.text = label
    run.append(text)
    hl.append(run)
    paragraph._p.append(hl)


doc = Document(DOC)
entries = {
    'Hussain Touhid Siddiquee is currently': {
        'scholar': 'https://scholar.google.com/citations?user=O8UVIQUAAAAJ&hl=en',
        'publons': 'https://www.webofscience.com/wos/author/record/QWB-6957-2026',
        'orcid': 'https://orcid.org/0009-0002-8804-6195',
    },
    'Orpon Chanda works as': {
        'scholar': 'https://scholar.google.com/citations?user=EYoY9JIAAAAJ',
        'publons': 'https://www.webofscience.com/wos/author/record/QWC-0602-2026',
        'orcid': 'https://orcid.org/0009-0004-3062-4989',
    },
}

for paragraph in doc.paragraphs:
    key = next((key for key in entries if paragraph.text.startswith(key)), None)
    if key is None:
        continue
    prose = paragraph.text.split(' Professional profiles:')[0]
    for child in list(paragraph._p):
        if child.tag != qn('w:pPr'):
            paragraph._p.remove(child)
    paragraph.add_run(prose + ' Professional profiles: Google Scholar: ')
    item = entries[key]
    link(paragraph, 'profile', item['scholar'])
    paragraph.add_run(' | Publons: ')
    if item['publons']:
        link(paragraph, 'profile', item['publons'])
    else:
        paragraph.add_run('link to be added')
    paragraph.add_run(' | ORCID: ')
    link(paragraph, item['orcid'].rsplit('/', 1)[-1], item['orcid'])

doc.save(DOC)
print('bio profile links updated')
