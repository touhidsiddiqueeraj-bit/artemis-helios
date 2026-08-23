"""
revision19.py — Insert A.3 multidimensional-sensitivity (V.A) + renumber tail
=============================================================================
1) Insert into Section V.A (before "B.  Cost-Benefit Analysis"):
   - boundedness/subordination + multidimensional-sensitivity paragraph
   - Fig. 13 heatmap image (inline, 3187700 EMU) + caption
2) Renumber figures moved one slot by the new Fig. 13:
   13 (BOM) -> 14, 14 (logger) -> 15, 15 (QR) -> 16, plus in-text refs.

Run:  python3 backups/revision19.py
"""
import re
import copy
from lxml import etree
from docx import Document
from docx.shared import Emu

DOC = '25195-52952-1-SM-REVISED.docx'
W_ = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
COL_W_EMU = 3187700
REMARK = {13: 14, 14: 15, 15: 16}
FIG = 'Figures/fig13_sensitivity_heatmap.png'

doc = Document(DOC)

def para_text(p):
    return ''.join(r.text or '' for r in p.runs)

def set_text(p, new):
    for r in p.runs:
        r.text = ''
    p.runs[0].text = new

# ---- 2) locate V.A anchors (BEFORE any in-memory text edits) ----
paras = list(doc.paragraphs)
cap10 = next(p for p in paras if para_text(p).strip().startswith('Fig. 10. Power-stage'))
cb   = next(p for p in paras if para_text(p).strip().startswith('B.  Cost-Benefit Analysis'))
body_tmpl = next(p for p in paras if para_text(p).strip().startswith('The power stage used in all simulations'))

# ---- 1) renumber tail figures (BOM 13->14, logger 14->15, QR 15->16) ----
FIGS_RE = re.compile(r'\b(Fig(?:ure)?s?\.)\s*(\d+)\s*([\u2013\u2014-]\s*(\d+))?\s*(\([a-d]\))?')
def sub(m):
    n1 = int(m.group(2)); out = f"{m.group(1)} {REMARK.get(n1, n1)}"
    if m.group(3):
        n2 = int(m.group(4)); sep = m.group(3).replace(str(n2), '')
        out += f"{sep} {REMARK.get(n2, n2)}"
    if m.group(5):
        out += m.group(5)
    return out

changed = 0
for p in doc.paragraphs:
    t = para_text(p)
    nt = FIGS_RE.sub(sub, t)
    if nt != t:
        set_text(p, nt)
        changed += 1
print('renumbered paragraphs:', changed)

# ---- 3) build elements ----
def clone_with_text(tmpl_el, text):
    el = copy.deepcopy(tmpl_el)
    runs = el.findall(f'{W_}r')
    first = runs[0]
    for r in runs[1:]:
        el.remove(r)
    ts = first.findall(f'{W_}t')
    if not ts:
        t = etree.SubElement(first, f'{W_}t')
        t.set(W_ + 'space', 'preserve')
        t.text = text
    else:
        for t in ts[1:]:
            first.remove(t)
        ts[0].text = text
    return el

def caption_el(text):
    return clone_with_text(cap10._p, text)

def image_el(path):
    tmp = doc.add_paragraph()
    pPr = etree.SubElement(tmp._p, f'{W_}pPr')
    jc = etree.SubElement(pPr, f'{W_}jc'); jc.set(W_ + 'val', 'center')
    tmp.add_run().add_picture(path, width=Emu(COL_W_EMU))
    return tmp._p

para = clone_with_text(body_tmpl._p,
 'The multidimensional sensitivity analysis extends the \u03b1 sweep by varying '
 'the blend deadband and the post-blend cooldown against \u03b1 (Fig. 13), with '
 'all other factors held at their Section III baseline (P&O gain k = 0.005, '
 'step limits [0.05, 0.80] V, AR(1) forecast, 100 ms update). Over the '
 '\u03b1 \u2208 [0.25, 0.45] \u00d7 deadband \u2208 [0.10, 0.20] region the tracking '
 'efficiency varies by less than 1 percentage point about the 95.1% baseline, '
 'and cooldowns of 10\u201320 control steps are near-optimal (93.5% at zero '
 'cooldown versus 95.1% at 20 steps), demonstrating that the reported '
 '\u03b1 = 0.35 / 15% deadband / 20-step settings sit inside a broad robustness '
 'plateau rather than a tuned optimum. Boundedness under prediction error '
 'follows directly from the blend law: the blended reference always remains '
 'within the 15% deviation deadband of the reactive P&O reference, so an '
 'erroneous forecast can displace the operating point by at most the deadband, '
 'after which the VS-P&O state machine recovers; the predictive term is '
 'therefore subordinate to the stabilizing reactive controller for arbitrary '
 'prediction errors.')

cap = caption_el(
 'Fig. 13. Multidimensional sensitivity of the LSTM-P&O tracking efficiency '
 'on the stochastic day (seed 23, 1 h, dt = 0.1 s): (a) \u03b1 \u00d7 deadband, (b) '
 '\u03b1 \u00d7 cooldown, all other factors at baseline. The \u03b1 \u2208 [0.25, 0.45], '
 'deadband \u2208 [10%, 20%] plateau varies by less than 1 pp about 95.1%).')

anchor = cb._p
for el in (para, image_el(FIG), cap):
    anchor.addprevious(el)
    anchor = el

doc.save(DOC)
print('revision19 applied: V.A sensitivity figure (Fig. 13) inserted, tail renumbered')
