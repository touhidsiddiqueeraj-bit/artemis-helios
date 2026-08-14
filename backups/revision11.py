"""
revision11.py — Insert new Section VII (HARDWARE IMPLEMENTATION) with
Figs. 12-13 (field logger schematic + PCB) into the revised manuscript.

Insertion point: immediately before the "Funding Information" paragraph
(after Section VI Conclusion). Figures are wrapped in full-width 1x1
borderless tables (dual-column template requirement, as in revision10.py).
Also appends a cross-reference to Section VII at the end of Section III.D.
"""
import re, shutil, copy, os
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Emu
from lxml import etree
assert etree is not None

DOCX = '25195-52952-1-SM-REVISED.docx'
BAK = '25195-52952-1-SM-REVISED-BAK13.docx'
FIG12 = 'Figures/fig12_hw_schematic.png'
FIG13 = 'Figures/fig13_hw_pcb.png'
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W_ = f'{{{W}}}'
FULL_W_TWIPS = 10466

if not os.path.exists(BAK):
    shutil.copy2(DOCX, BAK)
    print(f"Backup -> {BAK}")

doc = Document(DOCX)
body = doc.element.body
paras = list(body.iterchildren(f'{W_}p'))

def txt(el):
    return ''.join(t.text or '' for t in el.iter() if t.tag.endswith('}t')).strip()

def para_text(p):
    return txt(p)

# ── locate insertion point and template paragraphs ────────────────────────
fund = next(p for p in paras if para_text(p).startswith('Funding Information'))
concl = next(p for p in paras if para_text(p).startswith('VI.  CONCLUSION'))
body_tmpl = next(p for p in paras if para_text(p).startswith('The correction fires only'))
cap11 = next(p for p in paras if para_text(p).startswith('Fig. 11. '))
cross_tmpl = next(p for p in paras if para_text(p).startswith('The correction fires only'))
cross_target = next(p for p in paras if para_text(p).startswith('To provide the first empirical validation'))

def clone_with_text(tmpl_el, new_text):
    """Deep-copy paragraph, keep first run's formatting, set text, drop other runs."""
    el = copy.deepcopy(tmpl_el)
    runs = el.findall(f'{W_}r')
    assert runs, 'no runs in template para'
    first = runs[0]
    for r in runs[1:]:
        el.remove(r)
    ts = first.findall(f'{W_}t')
    if not ts:
        t = etree.SubElement(first, f'{W_}t')
        t.set(W_ + 'space', 'preserve')
        t.text = ''
    else:
        for t in ts[1:]:
            first.remove(t)
        ts[0].text = ''
    ts[0] if ts else None
    (first.findall(f'{W_}t')[0]).text = new_text
    return el

def make_tbl():
    tbl = etree.Element(f'{W_}tbl')
    tblPr = etree.SubElement(tbl, f'{W_}tblPr')
    tblW = etree.SubElement(tblPr, f'{W_}tblW')
    tblW.set(W_ + 'w', str(FULL_W_TWIPS)); tblW.set(W_ + 'type', 'dxa')
    etree.SubElement(tblPr, f'{W_}tblLayout').set(W_ + 'type', 'fixed')
    etree.SubElement(tblPr, f'{W_}jc').set(W_ + 'val', 'center')
    borders = etree.SubElement(tblPr, f'{W_}tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        etree.SubElement(borders, f'{W_}{side}').set(W_ + 'val', 'nil')
    mar = etree.SubElement(tblPr, f'{W_}tblCellMar')
    for side in ('top', 'left', 'bottom', 'right'):
        s = etree.SubElement(mar, f'{W_}{side}')
        s.set(W_ + 'w', '0'); s.set(W_ + 'type', 'dxa')
    grid = etree.SubElement(tbl, f'{W_}tblGrid')
    etree.SubElement(grid, f'{W_}gridCol').set(W_ + 'w', str(FULL_W_TWIPS))
    tr = etree.SubElement(tbl, f'{W_}tr')
    tc = etree.SubElement(tr, f'{W_}tc')
    tcPr = etree.SubElement(tc, f'{W_}tcPr')
    tcW = etree.SubElement(tcPr, f'{W_}tcW')
    tcW.set(W_ + 'w', str(FULL_W_TWIPS)); tcW.set(W_ + 'type', 'dxa')
    etree.SubElement(tcPr, f'{W_}tcMar')
    etree.SubElement(tcPr, f'{W_}vAlign').set(W_ + 'val', 'center')
    return tbl

def figure_table(img_path, width_in, caption_text, center=False):
    """Full-width table with image (python-docx picture) + caption paragraph."""
    tbl = make_tbl()
    tc = tbl.find(f'{W_}tr').find(f'{W_}tc')
    p_cap = etree.SubElement(tc, f'{W_}p')

    # image via python-docx run machinery
    tmp_para = doc.add_paragraph()
    run = tmp_para.add_run()
    run.add_picture(img_path, width=Emu(int(width_in * 914400)))
    if center:
        pPr = etree.SubElement(tmp_para._p, f'{W_}pPr')
        jc = etree.SubElement(pPr, f'{W_}jc'); jc.set(W_ + 'val', 'center')
    cap_el = clone_with_text(cap11, caption_text)
    tc.append(tmp_para._p)
    tc.append(cap_el)
    return tbl

# ── assemble new section elements in order ────────────────────────────────
heading = clone_with_text(concl, 'VII. HARDWARE IMPLEMENTATION')

p1 = clone_with_text(body_tmpl,
 'The field irradiance logger employed for the model validation reported in '
 'Sections III.D and IV.F was implemented as a purpose-built data-acquisition '
 'board (Fig. 12). The board carries the complete logger electronics on a '
 'compact PCB (roughly 92 mm × 80 mm, Fig. 13): an ESP32-S3 WROOM N16R8 CAM '
 'module (16 MB flash, 8 MB PSRAM) provides the logging, prediction, and '
 'communication core; a GY-302 BH1750 digital ambient-light sensor performs '
 'the irradiance measurements at the 10 s sampling interval behind the '
 'protective glass cover; a DS3231 real-time clock supplies absolute '
 'timestamps for the logged records; and an MP1584 buck converter derives '
 'the 3.3 V and 5 V logic rails from the 12 V SHS battery bus via the '
 'labelled battery terminals.')

p2 = clone_with_text(body_tmpl,
 'The PCB integrates the module pin headers, a decoupling capacitor bank '
 '(C2–C15) on the power rails, resistor networks, boot-configuration and '
 'reset switches, SD-card interfaces on GPIO38–GPIO40, and the UART link '
 '(GPIO43-TX/GPIO44-RX) that enables the Helios-to-Artemis communication '
 'protocol of Section III.A. The board layout follows the functional '
 'partitioning of the controller architecture: all logging, timestamping, '
 'and model-inference resources reside on the ESP32-S3 subsystem, with '
 'provision for the companion control plane through the UART interface. '
 'The assembled hardware was deployed for the Jul 9–14, 2026 field campaign '
 'described in Section III.D.')

tbl12 = figure_table(FIG12, 7.0,
 'Fig. 12. Field logger schematic: ESP32-S3 WROOM N16R8 CAM, GY-302 BH1750 '
 'light sensor, DS3231 RTC, MP1584 buck converter (12 V input → 3.3 V/5 V '
 'rails), and battery/switch interfaces (Leading University, REV 1.0).')

p3 = clone_with_text(body_tmpl,
 'Fig. 13 shows the resulting PCB layout. Silkscreen markers identify the '
 'functional blocks (Light Sensor, RTC Module, Boot Switch, 5 V, 3.3 V, GND) '
 'and the ±12 V battery terminals; four mounting holes and edge protection '
 'features support the weather-protected field enclosure. The sensor sits '
 'behind a clear glass cover whose attenuation was characterised in '
 'Section III.D (ratio 0.9314, factor 1.0737).')

tbl13 = figure_table(FIG13, 3.6,
 'Fig. 13. Field logger PCB layout: ESP32-S3 WROOM N16R8 CAM, BH1750 light '
 'sensor footprint, DS3231 RTC, MP1584 buck converter, capacitor bank '
 'C2–C15, boot switch, ±12 V battery terminals, mounting holes.', center=True)

# insert in order before the Funding Information paragraph
anchor = fund
for el in (heading, p1, tbl12, p2, p3, tbl13):
    anchor.addprevious(el)

# ── cross-reference in III.D ──────────────────────────────────────────────
cross_append = (' The logger hardware — schematic and PCB layout — is '
                'presented in Section VII, Figs. 12–13.')
first_run = cross_target.findall(f'{W_}r')[0]
rPr = first_run.find(f'{W_}rPr')
r = etree.SubElement(cross_target, f'{W_}r')
if rPr is not None:
    r.append(copy.deepcopy(rPr))
t = etree.SubElement(r, f'{W_}t')
t.set(W_ + 'space', 'preserve')
t.text = cross_append

doc.save(DOCX)
print('saved')

# ── verify ─────────────────────────────────────────────────────────────────
doc2 = Document(DOCX)
paras2 = doc2.element.body.findall('.//' + f'{W_}p')
blips = doc2.element.body.findall('.//' + qn('a:blip'))
tbls = doc2.element.body.findall(f'{W_}tbl')
print('blips total:', len(blips))
print('tables:', len(tbls))
print('section VII heading present:',
      any(para_text(p) == 'VII. HARDWARE IMPLEMENTATION' for p in paras2))
print('cross-ref present:',
      any('Figs. 12–13' in para_text(p) for p in paras2))
assert len(blips) == 13, 'expected 13 images'
assert any(para_text(p) == 'VII. HARDWARE IMPLEMENTATION' for p in paras2)
print('Section VII inserted OK')