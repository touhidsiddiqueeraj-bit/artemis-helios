"""
revision4.py — Text edits + figure re-embedding for final revision.

Changes:
1. Para 73: BOM cost — add SMT assembly 250 BDT, total 1,750 BDT, payback 6.1 yr
2. Para 75: Fig. 10 caption — update total
3. Para 87: Author Contributions — add "field deployment and hardware, O.C."
4. Para 89: Data Availability — replace "entirely simulation-based" with field validation statement
5. Re-embed figures: Fig 1, 6, 7, 8, 10 (5 figures)
"""
import os, re, shutil
from docx import Document
from docx.oxml.ns import qn

DOCX = '25195-52952-1-SM-REVISED.docx'
DOCX_BAK = '25195-52952-1-SM-REVISED-BAK4.docx'

shutil.copy2(DOCX, DOCX_BAK)
print(f"Backed up {DOCX} → {DOCX_BAK}")

doc = Document(DOCX)
paras = doc.paragraphs

# ── IMAGE REPLACEMENT ─────────────────────────────────────────────
# Map: paragraph index → (new image path, description)
fig_images = {
    13: ('Figures/figures_python/FIG1/output/fig1_architecture_fixed.png', 'Fig 1 architecture'),
    54: ('Figures/figures_python/FIG6/output/fig6_comparison.png',       'Fig 6 MPPT comparison'),
    59: ('Figures/figures_python/FIG7/output/fig7_po_convergence.png',   'Fig 7 P&O convergence'),
    63: ('Figures/figures_python/FIG9/output/fig9_validation.png',       'Fig 8 validation'),
    74: ('Figures/figures_python/FIG8/output/fig8_cost.png',             'Fig 10 BOM cost'),
}

rels = doc.part.rels
replaced_count = 0

for pi, (img_rel_path, label) in fig_images.items():
    img_abs = os.path.join(os.path.dirname(__file__) or '.', img_rel_path)
    if not os.path.exists(img_abs):
        print(f"  SKIP {label}: {img_abs} not found")
        continue

    p = paras[pi]
    blips = p._element.findall('.//' + qn('a:blip'))
    if not blips:
        print(f"  SKIP {label}: no blip in para {pi}")
        continue

    rid = blips[0].get(qn('r:embed'))
    rel = rels[rid]
    target_part = rel.target_part

    with open(img_abs, 'rb') as f:
        new_bytes = f.read()

    target_part._blob = new_bytes
    replaced_count += 1
    print(f"  Para {pi}: {label} ← {img_rel_path} ({len(new_bytes)} bytes)")

print(f"  Replaced {replaced_count}/{len(fig_images)} images")

# ── TEXT EDITS ────────────────────────────────────────────────────

def find_para(text_fragment):
    for pi, p in enumerate(paras):
        if text_fragment in p.text:
            return pi, p
    return None, None

def replace_para_text(pi, new_text):
    p = paras[pi]
    for run in p.runs:
        run._element.getparent().remove(run._element)
    run = p.add_run(new_text)
    print(f"  Para {pi}: updated ✓")

# 1. Para 73 — BOM cost: add assembly, update total/payback
pi73, _ = find_para('component cost is approximately')
if pi73 is not None:
    new73 = (
        'Fig. 10 presents the component cost breakdown. The estimated controller '
        'component cost (including SMT assembly) is approximately 1,750 BDT (USD 16), '
        'comprising ESP32-S3 (380 BDT), STM32F103 (120 BDT), INA219 (80 BDT), TSL2591 '
        '(120 BDT), buck converter passives (350 BDT), PCB and housing (280 BDT), '
        'miscellaneous connectors (170 BDT), and SMT assembly (250 BDT). This represents '
        'an 87% reduction versus IDCOL-compatible commercial MPPT controllers (13,500 BDT). '
        'At a rural avoided cost of 60 BDT/kWh, the projected annual yield improvement of '
        '+4.8 kWh per 50 Wp panel (91.3% vs 85.8% weighted annual efficiency, from Table III '
        'Monte Carlo annual means) generates +289 BDT/year additional energy value, implying '
        'a payback period of approximately 6.1 years against a plain P&O controller baseline.'
    )
    replace_para_text(pi73, new73)

# 2. Para 75 — Fig. 10 caption: update total
pi75, _ = find_para('Total ~1,500 BDT')
if pi75 is not None:
    new75 = (
        'Fig. 10.  BOM breakdown and cost comparison. Total ~1,750 BDT '
        '(including SMT assembly) — 87% below IDCOL-compatible commercial MPPT '
        '(Dhaka retail, Q1 2026).'
    )
    replace_para_text(pi75, new75)

# 3. Para 87 — Author Contributions: add field deployment credit
pi87, _ = find_para('field deployment')
if pi87 is not None:
    print(f"  Para {pi87}: already has field deployment credit — skipping")
else:
    pi87, _ = find_para('resources, H.T.S.')
    if pi87 is not None:
        t = paras[pi87].text
        new87 = t.replace('resources, H.T.S.;', 'resources, H.T.S.; field deployment and hardware, O.C.;')
        replace_para_text(pi87, new87)

# 4. Para 89 — Data Availability: replace "entirely simulation-based"
pi89, pa89 = find_para('entirely simulation-based')
if pi89 is not None:
    t = pa89.text
    new89 = t.replace(
        'This study is entirely simulation-based.',
        'The synthetic irradiance dataset was generated using a parametric Markov-chain + '
        'Ornstein-Uhlenbeck model parameterised from NASA POWER and SREDA publicly available '
        'data. The model is validated against field logger data (42 h daytime, G>80 W/m², '
        '1-minute resampled, Sylhet, Jul 10–13 2026).'
    )
    replace_para_text(pi89, new89)

# ── SAVE ──────────────────────────────────────────────────────────
doc.save(DOCX)
print(f"\nSaved {DOCX}")

# ── VERIFY ────────────────────────────────────────────────────────
doc2 = Document(DOCX)
print("\n=== VERIFICATION ===")
for search in ['1,750', '6.1 years', '1,500', 'resources, H.T.S.;', 'field deployment', 
               'entirely simulation', 'The synthetic irradiance', 'field logger data']:
    count = 0
    for pi, p in enumerate(doc2.paragraphs):
        if search in p.text:
            count += 1
            print(f"  [{search}] found in para {pi}: {p.text[:100]}")
    if count == 0:
        print(f"  [{search}] NOT FOUND — WARNING")
